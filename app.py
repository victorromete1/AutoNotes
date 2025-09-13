#imports
import io
import re
import calendar                     # For building the calendar view
import hashlib                      # For hashing passwords
from datetime import datetime       # For dates
from datetime import timedelta      # For duration 
from collections import defaultdict
import streamlit as st              
from PyPDF2 import PdfReader        # PDF text extraction
import docx                         # Docx text extraction

from supabase import create_client  # Supabase client

# Local Imports
import user_data
from note_generator import NoteGenerator
from flashcard_generator import FlashcardGenerator
from quiz_generator import QuizGenerator
from progress_tracker import ProgressTracker
from pdf_report_generator import PDFReportGenerator
from data_persistence import DataPersistence
from advanced_quiz_system import AdvancedQuizSystem
from data_import_export import DataImportExport
from utils import sanitize_filename
from autograder import AutoGrader




# Configuration & Setup

# App base configuration
st.set_page_config(
    page_title="AI Study Platform",
    # Stylish dashboard redesign
    else:
        username = st.session_state.get('username', 'User')
        st.markdown("""
        <style>
        .dashboard-container { max-width: 1000px; margin: 0 auto; }
        .dashboard-header { background: linear-gradient(90deg,#667eea,#764ba2); color: #fff; border-radius: 16px; padding: 28px 32px; margin-bottom: 24px; box-shadow: 0 6px 24px rgba(15,23,42,0.08); }
        .dashboard-title { font-size: 2.2rem; font-weight: 700; margin-bottom: 6px; }
        .dashboard-sub { font-size: 1.1rem; color: #e0e7ff; }
        .dashboard-stats { display: flex; gap: 18px; margin-bottom: 24px; }
        .dashboard-stat { background: #fff; color: #334155; border-radius: 12px; padding: 18px 0; flex: 1; text-align: center; box-shadow: 0 2px 8px rgba(15,23,42,0.04); }
        .dashboard-stat h3 { margin:0; font-size: 2rem; font-weight: 700; }
        .dashboard-stat p { margin:0; font-size: 1rem; color: #64748b; }
        .dashboard-main { display: flex; gap: 28px; }
        .dashboard-left { flex:2; }
        .dashboard-right { flex:1; }
        .dashboard-card { background: #fff; border-radius: 12px; padding: 20px 22px; margin-bottom: 20px; box-shadow: 0 4px 16px rgba(15,23,42,0.04); }
        .dashboard-actions { display: flex; gap: 12px; margin-bottom: 18px; }
        .dashboard-btn { background: linear-gradient(90deg,#667eea,#764ba2); color: #fff; border: none; padding: 12px 20px; border-radius: 8px; font-weight: 600; cursor: pointer; transition: box-shadow 0.2s; }
        .dashboard-btn:hover { box-shadow: 0 2px 12px rgba(102,126,234,0.15); }
        .dashboard-activity { background: #f8fafc; border-radius: 8px; padding: 14px 16px; margin-bottom: 10px; }
        .dashboard-muted { color: #64748b; font-size: 0.95rem; }
        </style>
        """, unsafe_allow_html=True)
        st.markdown(f"""
        <div class='dashboard-container'>
            <div class='dashboard-header'>
                <div class='dashboard-title'>&#127891; SmartStudy</div>
                <div class='dashboard-sub'>Welcome back, <strong>{username}</strong> &mdash; Your personal learning dashboard</div>
            </div>
        """, unsafe_allow_html=True)
        # Stats row
        stats = [
            (len(st.session_state.get('notes', [])), 'Notes'),
            (len(st.session_state.get('flashcards', [])), 'Flashcards'),
            (len([s for s in st.session_state.get('study_sessions', []) if s.get('activity_type') == 'quiz']), 'Quizzes Taken'),
            (len(st.session_state.get('study_sessions', [])), 'Study Sessions')
        ]
        st.markdown("<div class='dashboard-stats'>", unsafe_allow_html=True)
        for count, label in stats:
            st.markdown(f"<div class='dashboard-stat'><h3>{count}</h3><p>{label}</p></div>", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)
        # Main content columns
        left, right = st.columns([2,1])
        with left:
            # Focus Timer
            st.markdown("""
            <div class='dashboard-card'>
                <div style='font-weight:600; font-size:1.1rem; margin-bottom:8px;'>Focus Timer</div>
                <div style='display:flex; gap:10px; align-items:center;'>
                    <input id='focus_minutes' type='number' min='1' max='180' value='25' style='width:80px; padding:7px; border-radius:7px; border:1px solid #e2e8f0;' />
                    <button id='start_focus' class='dashboard-btn'>Start</button>
                    <button id='pause_focus' class='dashboard-btn' style='background:#94a3b8;'>Pause</button>
                    <button id='reset_focus' class='dashboard-btn' style='background:#e2e8f0; color:#0f172a;'>Reset</button>
                </div>
                <div style='margin-top:14px; font-size:2rem; font-weight:700;' id='focus_display'>25:00</div>
                <div class='dashboard-muted' style='margin-top:7px;'>Set a focus period and stay productive.</div>
                <script>
                let timer = null;
                let remaining = 25*60;
                const display = document.getElementById('focus_display');
                const startBtn = document.getElementById('start_focus');
                const pauseBtn = document.getElementById('pause_focus');
                const resetBtn = document.getElementById('reset_focus');
                const minutesInput = document.getElementById('focus_minutes');
                function formatTime(s){
                    const m = Math.floor(s/60); const r = s%60; return `${String(m).padStart(2,'0')}:${String(r).padStart(2,'0')}`;
                }
                function tick(){
                    if(remaining<=0){ clearInterval(timer); timer=null; display.innerText='00:00'; return; }
                    remaining -=1; display.innerText = formatTime(remaining);
                }
                startBtn.onclick = ()=>{
                    remaining = parseInt(minutesInput.value||25)*60; display.innerText = formatTime(remaining);
                    if(timer) clearInterval(timer);
                    timer = setInterval(tick,1000);
                };
                pauseBtn.onclick = ()=>{ if(timer){ clearInterval(timer); timer=null;} else { timer = setInterval(tick,1000); } };
                resetBtn.onclick = ()=>{ if(timer){ clearInterval(timer); timer=null;} remaining = parseInt(minutesInput.value||25)*60; display.innerText = formatTime(remaining); };
                </script>
            </div>
            """, unsafe_allow_html=True)
            # Quick Actions
            st.markdown("<div class='dashboard-card'>", unsafe_allow_html=True)
            st.markdown("<div style='font-weight:600; font-size:1.1rem; margin-bottom:8px;'>Quick Actions</div>", unsafe_allow_html=True)
            actions = [
                ("New Note", "📝 Notes", "quick_note"),
                ("Study Flashcards", "📚 Flashcards", "quick_flashcards"),
                ("Take Quiz", "🧠 Quizzes", "quick_quiz")
            ]
            ac1, ac2, ac3 = st.columns(3)
            for (label, page, key), col in zip(actions, [ac1, ac2, ac3]):
                with col:
                    if st.button(label, use_container_width=True, key=key):
                        st.session_state.page = page
                        st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)
            # Recent Activity
            st.markdown("<div class='dashboard-card'>", unsafe_allow_html=True)
            st.markdown("<div style='font-weight:600; font-size:1.1rem; margin-bottom:8px;'>Recent Activity</div>", unsafe_allow_html=True)
            if st.session_state.get('study_sessions'):
                recent_sessions = sorted(st.session_state.get('study_sessions', []), key=lambda x: x.get('timestamp',''), reverse=True)[:6]
                for s in recent_sessions:
                    try:
                        ts = datetime.fromisoformat(s.get('timestamp')).strftime("%b %d %Y %H:%M")
                    except:
                        ts = s.get('timestamp','')
                    atype = s.get('activity_type','activity')
                    if atype == 'quiz':
                        score = s.get('score',0)
                        color = '#2ca02c' if score>=80 else '#ff7f0e' if score>=60 else '#d62728'
                        st.markdown(f"<div class='dashboard-activity'><div style='display:flex; justify-content:space-between; align-items:center;'><div><strong>Quiz</strong> • {s.get('subject','General')}</div><div style='color:{color}; font-weight:600'>{score:.1f}%</div></div><div class='dashboard-muted'>{ts}</div></div>", unsafe_allow_html=True)
                    elif atype == 'flashcards':
                        studied = s.get('flashcards_studied',0)
                        correct = s.get('correct_answers',0)
                        acc = (correct/studied*100) if studied>0 else 0
                        st.markdown(f"<div class='dashboard-activity'><div style='display:flex; justify-content:space-between; align-items:center;'><div><strong>Flashcards</strong> • {s.get('subject','General')}</div><div style='font-weight:600'>{studied} cards</div></div><div class='dashboard-muted'>{ts} • Accuracy: {acc:.1f}%</div></div>", unsafe_allow_html=True)
                    else:
                        st.markdown(f"<div class='dashboard-activity'><div style='display:flex; justify-content:space-between; align-items:center;'><div><strong>{atype.capitalize()}</strong></div><div class='dashboard-muted'>{ts}</div></div></div>", unsafe_allow_html=True)
            else:
                st.info("No recent activity. Start studying to see your progress here!")
            st.markdown("</div>", unsafe_allow_html=True)
            # Quick Note
            st.markdown("<div class='dashboard-card'>", unsafe_allow_html=True)
            st.markdown("<div style='font-weight:600; font-size:1.1rem; margin-bottom:8px;'>Quick Note</div>", unsafe_allow_html=True)
            with st.form("quick_note_form"):
                quick_note = st.text_area("Jot something down:", placeholder="Type your quick note here...", height=120, label_visibility="collapsed", key="quick_note_text")
                if st.form_submit_button("Save Quick Note", use_container_width=True):
                    if quick_note.strip():
                        new_note = {"title": f"Quick Note - {datetime.now().strftime('%H:%M')}", "content": quick_note, "category": "Quick Notes", "timestamp": datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
                        st.session_state.notes.append(new_note)
                        auto_save()
                        st.success("Quick note saved!")
                        st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)
        with right:
            # Upcoming Events
            st.markdown("<div class='dashboard-card'>", unsafe_allow_html=True)
            st.markdown("<div style='font-weight:600; font-size:1.1rem; margin-bottom:8px;'>Upcoming Events</div>", unsafe_allow_html=True)
            if st.session_state.get('events'):
                today = datetime.now().date()
                upcoming = []
                for ev in st.session_state.get('events',[]):
                    try:
                        d = datetime.fromisoformat(ev.get('date')).date()
                        if d >= today:
                            upcoming.append({'name':ev.get('name'), 'date':d, 'color': ev.get('color','#4CAF50')})
                    except:
                        continue
                upcoming = sorted(upcoming, key=lambda x: x['date'])[:3]
                for ev in upcoming:
                    days = (ev['date'] - today).days
                    when = 'Today' if days==0 else f'In {days} day' + ('s' if days!=1 else '')
                    st.markdown(f"<div style='background:#f9fafb; padding:10px; border-radius:8px; margin-bottom:8px; border-left:4px solid {ev['color']}'> <div style='font-weight:600'>{ev['name']}</div> <div class='dashboard-muted'>{ev['date'].strftime('%b %d')} • {when}</div> </div>", unsafe_allow_html=True)
            else:
                st.markdown("<div class='dashboard-muted'>No upcoming events.</div>", unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)
            # Study Tip
            st.markdown("<div class='dashboard-card'>", unsafe_allow_html=True)
            st.markdown("<div style='font-weight:600; font-size:1.1rem; margin-bottom:8px;'>Study Tip</div>", unsafe_allow_html=True)
            import random
            tips = [
                "Try the Pomodoro technique: 25 minutes focused, 5 minutes break.",
                "Teach someone else to reinforce learning.",
                "Create flashcards for key concepts and review regularly.",
            ]
            st.info(random.choice(tips))
            st.markdown("</div>", unsafe_allow_html=True)
            # Progress mini-chart
            st.markdown("<div class='dashboard-card'>", unsafe_allow_html=True)
            st.markdown("<div style='font-weight:600; font-size:1.1rem; margin-bottom:8px;'>This Week's Progress</div>", unsafe_allow_html=True)
            if st.session_state.get('study_sessions'):
                daily_counts = {}
                for session in st.session_state.get('study_sessions',[]):
                    try:
                        sd = datetime.fromisoformat(session.get('timestamp')).date()
                        if (datetime.now().date() - sd).days <= 7:
                            daily_counts[sd] = daily_counts.get(sd,0) + 1
                    except:
                        continue
                if daily_counts:
                    max_count = max(daily_counts.values())
                    dates = [(datetime.now().date() - timedelta(days=i)) for i in range(6, -1, -1)]
                    bars = "<div style='display:flex; gap:6px; align-items:flex-end; height:90px;'>"
                    for d in dates:
                        cnt = daily_counts.get(d,0)
                        h = int((cnt/max_count)*70) if max_count>0 else 0
                        bars += f"<div style='text-align:center; font-size:11px;'><div style='background:linear-gradient(90deg,#667eea,#764ba2); width:18px; height:{h}px; border-radius:4px;'></div><div style='margin-top:6px;'>{d.strftime('%a')}</div></div>"
                    bars += "</div>"
                    st.markdown(bars, unsafe_allow_html=True)
                else:
                    st.markdown("<div class='dashboard-muted'>No study sessions this week yet.</div>", unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)
        # Admin Controls
        if st.session_state.get("admin_mode"):
            st.markdown("---")
            st.markdown("<div class='dashboard-card'>", unsafe_allow_html=True)
            st.markdown("<h3>Admin Controls</h3>", unsafe_allow_html=True)
            admin_col1, admin_col2 = st.columns(2)
            with admin_col1:
                st.markdown("#### Reset Password")
                target_user = st.text_input("Username to Reset Password", key="admin_reset_user")
                new_pass = st.text_input("New Password", type="password", key="admin_new_pass")
                if st.button("Reset User Password", key="admin_reset_btn"):
                    if target_user and new_pass:
                        ok, msg = user_data.admin_reset_password(target_user, new_pass)
                        if ok:
                            st.success(f"Password for '{target_user}' reset successfully.")
                        else:
                            st.error(msg)
                    else:
                        st.warning("Enter both username and new password.")
            with admin_col2:
                st.markdown("#### Delete Account")
                del_user = st.text_input("Username to Delete", key="admin_del_user")
                if st.button("Delete Account", key="admin_del_btn"):
                    if del_user:
                        ok, msg = user_data.admin_delete_account(del_user)
                        if ok:
                            st.success(f"User '{del_user}' deleted successfully.")
                        else:
                            st.error(msg)
                    else:
                        st.warning("Enter a username to delete.")
            st.markdown("</div>", unsafe_allow_html=True)
            st.markdown("""
            ### Transform Your Learning Experience
            
            SmartStudy is your **all-in-one learning companion** designed to help you study smarter, not harder.
            
            ##### 🌟 Key Features:
            - **📝 Smart Notes** - AI-powered note generation and summarization
            - **🎴 Flashcards** - Create and study with interactive flashcards
            - **🧠 Adaptive Quizzes** - Test your knowledge with AI-generated quizzes
            - **📊 Progress Tracking** - Monitor your learning journey with detailed analytics
            - **📅 Calendar Integration** - Schedule study sessions and set reminders
            - **📝 Autograder** - Get instant feedback on your writing
            
            ##### 🚀 Get Started:
            1. Create an account or log in using the sidebar
            2. Explore the different study tools
            3. Track your progress over time
            4. Achieve your learning goals faster!
            """)
        
        with col2:
            st.image("https://images.unsplash.com/photo-1501504905252-473c47e087f8?ixlib=rb-1.2.1&auto=format&fit=crop&w=500&q=80", 
                    caption="Study Smarter with AI")
            
        # Feature cards
        st.markdown("### ✨ How It Works")
        col1, col2, col3 = st.columns(3)

        with col1:
            st.markdown("""
            <div class="feature-card" style="color: black;">
                <h4>&#128221; Create Content</h4>
                <ul style='padding-left: 1em;'>
                    <li>Generate notes from your materials</li>
                    <li>Create flashcards from your content</li>
                    <li>AI assistance for summarization</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)
        with col2:
            st.markdown("""
            <div class="feature-card" style="color: black;">
                <h4>&#129504; Study Smart</h4>
                <ul style='padding-left: 1em;'>
                    <li>Spaced repetition</li>
                    <li>Adaptive quizzes</li>
                    <li>Reinforce learning effectively</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)
        with col3:
            st.markdown("""
            <div class="feature-card" style="color: black;">
                <h4>&#128200; Track Progress</h4>
                <ul style='padding-left: 1em;'>
                    <li>Monitor your performance</li>
                    <li>Detailed analytics</li>
                    <li>Identify areas for improvement</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)

    # Logged in → Notion-like dashboard
    else:
        username = st.session_state.get('username', 'User')

        # Notion-like CSS
        st.markdown("""
        <style>
        .notion-container { max-width: 1100px; margin: 0 auto 24px; }
        .notion-header { background: #ffffff; border-radius: 12px; padding: 18px 22px; box-shadow: 0 6px 18px rgba(15, 23, 42, 0.06); margin-bottom: 18px; display:flex; justify-content:space-between; align-items:center; }
        .notion-title { font-size: 1.5rem; font-weight: 600; color: #0f172a; }
        .notion-sub { color: #475569; font-size: 0.95rem; }
        .notion-stats { display:flex; gap:12px; margin-bottom:18px; }
        .notion-stat { background:#fff; padding:14px; border-radius:10px; box-shadow: 0 4px 12px rgba(2,6,23,0.04); min-width:140px; text-align:center; }
        .notion-stat h3 { margin:0; font-size:1.25rem; color:#0f172a; }
        .notion-stat p { margin:0; color:#64748b; font-size:0.85rem; }
        .notion-main { display:flex; gap:20px; }
        .notion-left { flex:2; }
        .notion-right { flex:1; }
        .notion-card { background:#fff; border-radius:10px; padding:16px; box-shadow: 0 6px 18px rgba(15,23,42,0.04); margin-bottom:16px; }
        .quick-actions { display:flex; gap:10px; }
        .qa-btn { background:linear-gradient(90deg,#667eea,#764ba2); color:white; border:none; padding:10px 14px; border-radius:8px; cursor:pointer; }
        .activity-item { background:#f8fafc; padding:12px; border-radius:8px; margin-bottom:10px; }
        .small-muted { color:#64748b; font-size:0.9rem; }
        </style>
        """, unsafe_allow_html=True)

        # Header
        st.markdown(f"""
        <div class="notion-container">
            <div class="notion-header">
                <div>
                    <div class="notion-title">&#127891; SmartStudy</div>
                    <div class="notion-sub">Welcome back, <strong>{username}</strong></div>
                </div>
                <div class="notion-sub">Your personal learning workspace</div>
            </div>
        """, unsafe_allow_html=True)

        # Stats row
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.markdown(f"""
            <div class="notion-stat">
                <h3>{len(st.session_state.get('notes', []))}</h3>
                <p>Notes</p>
            </div>
            """, unsafe_allow_html=True)
        with col2:
            st.markdown(f"""
            <div class="notion-stat">
                <h3>{len(st.session_state.get('flashcards', []))}</h3>
                <p>Flashcards</p>
            </div>
            """, unsafe_allow_html=True)
        with col3:
            quiz_sessions = [s for s in st.session_state.get('study_sessions', []) if s.get('activity_type') == 'quiz']
            st.markdown(f"""
            <div class="notion-stat">
                <h3>{len(quiz_sessions)}</h3>
                <p>Quizzes Taken</p>
            </div>
            """, unsafe_allow_html=True)
        with col4:
            st.markdown(f"""
            <div class="notion-stat">
                <h3>{len(st.session_state.get('study_sessions', []))}</h3>
                <p>Study Sessions</p>
            </div>
            """, unsafe_allow_html=True)

        # Main content columns
        left_col, right_col = st.columns([2,1])

        with left_col:
            # Quick actions (Streamlit buttons for interactivity)
            # Focus timer UI
            st.markdown("""
            <div class="notion-card">
                <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom:8px;">
                    <div style="font-weight:600;">Focus Timer</div>
                    <div class="small-muted">Stay focused</div>
                </div>
                <div style='display:flex; gap:8px; align-items:center;'>
                    <input id='focus_minutes' type='number' min='1' max='180' value='25' style='width:80px; padding:6px; border-radius:6px; border:1px solid #e2e8f0;' />
                    <button id='start_focus' class='qa-btn'>Start</button>
                    <button id='pause_focus' class='qa-btn' style='background:#94a3b8;'>Pause</button>
                    <button id='reset_focus' class='qa-btn' style='background:#e2e8f0; color:#0f172a;'>Reset</button>
                </div>
                <div style='margin-top:12px; font-size:1.5rem; font-weight:600;' id='focus_display'>25:00</div>
                <div class='small-muted' style='margin-top:6px;'>Use this timer to set a focused study period.</div>
                <script>
                // Simple client-side timer
                let timer = null;
                let remaining = 25*60;
                const display = document.getElementById('focus_display');
                const startBtn = document.getElementById('start_focus');
                const pauseBtn = document.getElementById('pause_focus');
                const resetBtn = document.getElementById('reset_focus');
                const minutesInput = document.getElementById('focus_minutes');

                function formatTime(s){
                    const m = Math.floor(s/60); const r = s%60; return `${String(m).padStart(2,'0')}:${String(r).padStart(2,'0')}`;
                }
                function tick(){
                    if(remaining<=0){ clearInterval(timer); timer=null; display.innerText='00:00'; return; }
                    remaining -=1; display.innerText = formatTime(remaining);
                }
                startBtn.onclick = ()=>{
                    remaining = parseInt(minutesInput.value||25)*60; display.innerText = formatTime(remaining);
                    if(timer) clearInterval(timer);
                    timer = setInterval(tick,1000);
                };
                pauseBtn.onclick = ()=>{ if(timer){ clearInterval(timer); timer=null;} else { timer = setInterval(tick,1000); } };
                resetBtn.onclick = ()=>{ if(timer){ clearInterval(timer); timer=null;} remaining = parseInt(minutesInput.value||25)*60; display.innerText = formatTime(remaining); };
                </script>
            """, unsafe_allow_html=True)

            qa1, qa2, qa3 = st.columns(3)
            with qa1:
                if st.button("New Note", use_container_width=True, key="quick_note"):
                    st.session_state.page = "📝 Notes"
                    st.rerun()
            with qa2:
                if st.button("Study Flashcards", use_container_width=True, key="quick_flashcards"):
                    st.session_state.page = "📚 Flashcards"
                    st.rerun()
            with qa3:
                if st.button("Take Quiz", use_container_width=True, key="quick_quiz"):
                    st.session_state.page = "🧠 Quizzes"
                    st.rerun()

            # Recent Activity
            st.markdown("""
                </div>
            """, unsafe_allow_html=True)

            st.markdown("""
            <div class="notion-card">
                <div style="font-weight:600; margin-bottom:8px;">Recent Activity</div>
            """, unsafe_allow_html=True)

            if st.session_state.get('study_sessions'):
                recent_sessions = sorted(st.session_state.get('study_sessions', []), key=lambda x: x.get('timestamp',''), reverse=True)[:6]
                for s in recent_sessions:
                    try:
                        ts = datetime.fromisoformat(s.get('timestamp')).strftime("%b %d %Y %H:%M")
                    except:
                        ts = s.get('timestamp','')
                    atype = s.get('activity_type','activity')
                    if atype == 'quiz':
                        score = s.get('score',0)
                        color = '#2ca02c' if score>=80 else '#ff7f0e' if score>=60 else '#d62728'
                        st.markdown(f"""
                        <div class='activity-item'>
                            <div style='display:flex; justify-content:space-between; align-items:center;'>
                                <div><strong>Quiz</strong> • {s.get('subject','General')}</div>
                                <div style='color:{color}; font-weight:600'>{score:.1f}%</div>
                            </div>
                            <div class='small-muted'>{ts}</div>
                        </div>
                        """, unsafe_allow_html=True)
                    elif atype == 'flashcards':
                        studied = s.get('flashcards_studied',0)
                        correct = s.get('correct_answers',0)
                        acc = (correct/studied*100) if studied>0 else 0
                        st.markdown(f"""
                        <div class='activity-item'>
                            <div style='display:flex; justify-content:space-between; align-items:center;'>
                                <div><strong>Flashcards</strong> • {s.get('subject','General')}</div>
                                <div style='font-weight:600'>{studied} cards</div>
                            </div>
                            <div class='small-muted'>{ts} • Accuracy: {acc:.1f}%</div>
                        </div>
                        """, unsafe_allow_html=True)
                    else:
                        st.markdown(f"""
                        <div class='activity-item'>
                            <div style='display:flex; justify-content:space-between; align-items:center;'>
                                <div><strong>{atype.capitalize()}</strong></div>
                                <div class='small-muted'>{ts}</div>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
            else:
                st.info("No recent activity. Start studying to see your progress here!")

            # Quick Note form
            st.markdown("""
            </div>
            """, unsafe_allow_html=True)

            with st.form("quick_note_form"):
                quick_note = st.text_area("Jot something down:", placeholder="Type your quick note here...", height=120, label_visibility="collapsed", key="quick_note_text")
                if st.form_submit_button("Save Quick Note", use_container_width=True):
                    if quick_note.strip():
                        new_note = {"title": f"Quick Note - {datetime.now().strftime('%H:%M')}", "content": quick_note, "category": "Quick Notes", "timestamp": datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
                        st.session_state.notes.append(new_note)
                        auto_save()
                        st.success("Quick note saved!")
                        st.rerun()

        with right_col:
            # Upcoming events
            st.markdown("""
            <div class='notion-card'>
                <div style='font-weight:600; margin-bottom:8px;'>Upcoming Events</div>
            """, unsafe_allow_html=True)
            if st.session_state.get('events'):
                today = datetime.now().date()
                upcoming = []
                for ev in st.session_state.get('events',[]):
                    try:
                        d = datetime.fromisoformat(ev.get('date')).date()
                        if d >= today:
                            upcoming.append({'name':ev.get('name'), 'date':d, 'color': ev.get('color','#4CAF50')})
                    except:
                        continue
                upcoming = sorted(upcoming, key=lambda x: x['date'])[:3]
                for ev in upcoming:
                    days = (ev['date'] - today).days
                    when = 'Today' if days==0 else f'In {days} day' + ('s' if days!=1 else '')
                    st.markdown(f"""
                    <div style='background:#f9fafb; padding:10px; border-radius:8px; margin-bottom:8px; border-left:4px solid {ev['color']}'>
                        <div style='font-weight:600'>{ev['name']}</div>
                        <div class='small-muted'>{ev['date'].strftime('%b %d')} • {when}</div>
                    </div>
                    """, unsafe_allow_html=True)
            else:
                st.markdown("<div class='small-muted'>No upcoming events.</div>", unsafe_allow_html=True)

            # Study tip
            st.markdown("""
            </div>
            <div class='notion-card'>
                <div style='font-weight:600; margin-bottom:8px;'>Study Tip</div>
            """, unsafe_allow_html=True)
            import random
            tips = [
                "Try the Pomodoro technique: 25 minutes focused, 5 minutes break.",
                "Teach someone else to reinforce learning.",
                "Create flashcards for key concepts and review regularly.",
            ]
            st.info(random.choice(tips))

            # Progress mini-chart (reuse existing logic)
            if st.session_state.get('study_sessions'):
                daily_counts = {}
                for session in st.session_state.get('study_sessions',[]):
                    try:
                        sd = datetime.fromisoformat(session.get('timestamp')).date()
                        if (datetime.now().date() - sd).days <= 7:
                            daily_counts[sd] = daily_counts.get(sd,0) + 1
                    except:
                        continue
                if daily_counts:
                    max_count = max(daily_counts.values())
                    dates = [(datetime.now().date() - timedelta(days=i)) for i in range(6, -1, -1)]
                    bars = "<div style='display:flex; gap:6px; align-items:flex-end; height:90px;'>"
                    for d in dates:
                        cnt = daily_counts.get(d,0)
                        h = int((cnt/max_count)*70) if max_count>0 else 0
                        bars += f"<div style='text-align:center; font-size:11px;'><div style='background:linear-gradient(90deg,#667eea,#764ba2); width:18px; height:{h}px; border-radius:4px;'></div><div style='margin-top:6px;'>{d.strftime('%a')}</div></div>"
                    bars += "</div>"
                    st.markdown(bars, unsafe_allow_html=True)
                else:
                    st.markdown("<div class='small-muted'>No study sessions this week yet.</div>", unsafe_allow_html=True)

        # Close container
        st.markdown("""
        </div>
        """, unsafe_allow_html=True)

        # Admin Controls preserved
        if st.session_state.get("admin_mode"):
            st.markdown("---")
            st.markdown("""
            <div class="notion-card">
                <h3>Admin Controls</h3>
            """, unsafe_allow_html=True)
            admin_col1, admin_col2 = st.columns(2)
            with admin_col1:
                st.markdown("#### Reset Password")
                target_user = st.text_input("Username to Reset Password", key="admin_reset_user")
                new_pass = st.text_input("New Password", type="password", key="admin_new_pass")
                if st.button("Reset User Password", key="admin_reset_btn"):
                    if target_user and new_pass:
                        ok, msg = user_data.admin_reset_password(target_user, new_pass)
                        if ok:
                            st.success(f"Password for '{target_user}' reset successfully.")
                        else:
                            st.error(msg)
                    else:
                        st.warning("Enter both username and new password.")
            with admin_col2:
                st.markdown("#### Delete Account")
                del_user = st.text_input("Username to Delete", key="admin_del_user")
                if st.button("Delete Account", key="admin_del_btn"):
                    if del_user:
                        ok, msg = user_data.admin_delete_account(del_user)
                        if ok:
                            st.success(f"User '{del_user}' deleted successfully.")
                        else:
                            st.error(msg)
                    else:
                        st.warning("Enter a username to delete.")
            st.markdown("</div>", unsafe_allow_html=True)


# ============================
# Notes Page
# ============================


elif st.session_state.page == "📝 Notes":
    st.title("📝 AI Note Generator & Class Notes")

    # Ensure keys used by this page exist
    for key in ["free_note_text", "free_note_title", "summarize_option", "free_note_cat", "topic_input", "topic_cat_input"]:
        if key not in st.session_state:
            if "text" in key or "title" in key or "input" in key:
                st.session_state[key] = ""
            elif "cat" in key:
                st.session_state[key] = "General"
            elif "summarize_option" in key:
                st.session_state[key] = True

    # Freeform Notes Mode
    st.subheader("💡 Freeform Notes Mode")

    note_title = st.text_input(
        "Note Name:",
        placeholder="Enter a title for your note...",
        key="free_note_title"
    )

    free_note = st.text_area(
        "Type your notes here:",
        value="",
        placeholder="Write your class notes here...",
        height=200,
        key="free_note_text"
    )

    free_category = st.text_input(
        "Category for these notes:",
        value=st.session_state.get("free_note_cat", "General"),
        key="free_note_cat"
    )

    summarize_option = st.checkbox(
        "🧠 Summarize with AI",
        value=st.session_state.get("summarize_option", True),
        key="summarize_option"
    )

    if st.button("💾 Save Notes", key="save_free_note"):
        user_data.save_current_user(st.session_state)
        if not free_note.strip():
            st.warning("Please enter some text to save.")
        else:
            notes_content = free_note
            if summarize_option:
                # AI summarize using NoteGenerator
                with st.spinner("Summarizing notes with AI..."):
                    try:
                        notes_content = generators['notes'].generate_notes(notes_content)
                    except Exception as e:
                        st.error(f"AI summarization failed: {e}")

            final_title = (note_title or "").strip() or "Untitled Note"

            # Append new note
            new_note = {
                "title": final_title,
                "content": notes_content,
                "category": free_category or "General",
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            st.session_state.notes.append(new_note)
            user_data.save_current_user(st.session_state)
            auto_save()
            st.success(f"✅ Note '{final_title}' saved!")
            st.rerun()

    st.markdown("---")

    # AI Note Generation from Topic or File
    st.subheader("🚀 Generate AI Notes from Topic or File")
    col1, col2 = st.columns([2, 1])
    with col1:
        topic = st.text_input("📖 Topic or subject:", placeholder="e.g., Photosynthesis, WWII, Calculus...", key="topic_input")
    with col2:
        category = st.text_input("Category:", value=st.session_state.topic_cat_input, key="topic_cat_input")

    uploaded_file = st.file_uploader(
        "📂 Upload a file (.txt, .md, .pdf, .docx) to generate notes:",
        type=['txt', 'md', 'pdf', 'docx'],
        key="file_upload"
    )

    if st.button("🚀 Generate Notes", key="generate_ai_notes"):
        user_data.save_current_user(st.session_state)
        content_to_process = ""
        note_name = ""

        # Extract content from uploaded file
        if uploaded_file:
            file_ext = uploaded_file.name.split('.')[-1].lower()
            note_name = uploaded_file.name.rsplit('.', 1)[0]
            try:
                if file_ext in ['txt', 'md']:
                    content_to_process = uploaded_file.read().decode("utf-8")
                elif file_ext == 'pdf':
                    pdf = PdfReader(io.BytesIO(uploaded_file.read()))
                    content_to_process = "\n".join([page.extract_text() or "" for page in pdf.pages])
                elif file_ext == 'docx':
                    doc = docx.Document(io.BytesIO(uploaded_file.read()))
                    content_to_process = "\n".join([para.text for para in doc.paragraphs])
                else:
                    st.error("Unsupported file type.")
            except Exception as e:
                st.error(f"Failed to read file: {e}")
        elif topic.strip():
            # Use typed topic directly
            content_to_process = topic
            note_name = topic
            

        # Generate notes via NoteGenerator
        if content_to_process:
            with st.spinner("Generating comprehensive notes..."):
                try:
                    notes_content = generators['notes'].generate_notes(content_to_process)
                    if notes_content:
                        new_note = {
                            "title": note_name or f"Note {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                            "content": notes_content,
                            "category": category or "General",
                            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        }
                        st.session_state.notes.append(new_note)
                        auto_save()
                        st.success(f"✅ Notes generated successfully for '{note_name}'!")
                        with st.expander("📖 Preview Generated Notes", expanded=True):
                            st.markdown(notes_content)
                            user_data.save_current_user(st.session_state)
                    else:
                        st.error("Failed to generate notes. Please try again.")
                except Exception as e:
                    st.error(f"Error: {e}")
        else:
            st.warning("Please enter a topic or upload a file.")
    # Existing Notes List
    if st.session_state.notes:
        st.divider()
        st.subheader("📚 Your Notes")

        categories = list(set([note.get('category', 'General') for note in st.session_state.notes]))
        filter_category = st.selectbox("Filter by category:", ["All"] + categories)

        filtered_notes = st.session_state.notes
        if filter_category != "All":
            filtered_notes = [n for n in filtered_notes if n.get('category') == filter_category]

        for i, note in enumerate(filtered_notes):
            with st.expander(f"📄 {note['title']} ({note.get('category', 'General')})"):
                st.write(f"**Created:** {note['timestamp']}")
                st.markdown(note['content'])

                col1, col2, col3, col4 = st.columns([1, 1, 1, 1])

                # Create flashcards from this note
                with col1:
                    if st.button("📚 Create Flashcards", key=f"flash_{i}"):
                        user_data.save_current_user(st.session_state)
                        with st.spinner("Creating flashcards..."):
                            try:
                                flashcards = generators['flashcards'].generate_flashcards(
                                    note['content'], num_cards=6, difficulty="Medium"
                                )
                                for card in flashcards:
                                    card['category'] = note.get('category', 'General')
                                st.session_state.flashcards.extend(flashcards)
                                auto_save()
                                st.success(f"✅ Created {len(flashcards)} flashcards!")

                                # Log activity
                                session = {
                                    "timestamp": datetime.now().isoformat(),
                                    "activity_type": "flashcards_created",
                                    "subject": note.get('category', 'General'),
                                    "flashcards_created": len(flashcards),
                                    "duration_minutes": 2
                                }
                                st.session_state.study_sessions.append(session)
                                auto_save()
                            except Exception as e:
                                st.error(f"Error: {e}")

                # Download note as .txt
                with col2:
                    st.download_button(
                        "📥 Download",
                        data=note['content'],
                        file_name=f"{sanitize_filename(note['title'])}.txt",
                        mime="text/plain",
                        key=f"download_{i}"
                    )

                # Delete note
                with col3:
                    if st.button("🗑️ Delete", key=f"delete_{i}"):
                        st.session_state.notes.remove(note)
                        user_data.save_current_user(st.session_state)
                        auto_save()
                        st.rerun()

                # Rename note
                with col4:
                    new_name = st.text_input("Rename Note", value=note['title'], key=f"rename_{i}")
                    if st.button("✏️ Rename", key=f"rename_btn_{i}"):
                        user_data.save_current_user(st.session_state)
                        if new_name.strip():
                            note['title'] = new_name.strip()
                            auto_save()
                            st.success("✅ Note renamed!")
                            st.rerun()
                        else:
                            st.warning("Enter a valid name.")


# ============================
# Flashcards Page
# ============================

elif st.session_state.page == "📚 Flashcards":
    st.title("📚 Interactive Flashcards")

    tab1, tab2, tab3 = st.tabs(["📖 Study", "➕ Create", "📂 Manage"])

    # Study Tab
    with tab1:
        st.subheader("📖 Study Session")

        if not st.session_state.flashcards:
            st.info("No flashcards available. Create some first!")
            if st.button("🔄 Refresh"):
                user_data.save_current_user(st.session_state)
                st.rerun()
        else:
            # Filter by category
            categories = list(set([card.get('category', 'General') for card in st.session_state.flashcards]))
            selected_category = st.selectbox("Study category:", ["All"] + categories)

            study_cards = st.session_state.flashcards
            if selected_category != "All":
                study_cards = [card for card in study_cards if card.get('category', 'General') == selected_category]

            if study_cards:
                # Initialize study state keys
                if 'study_index' not in st.session_state:
                    st.session_state.study_index = 0
                    st.session_state.show_answer = False
                    st.session_state.cards_studied = 0
                    st.session_state.cards_correct = 0

                current_card = study_cards[st.session_state.study_index]

                # Progress bar
                progress = (st.session_state.study_index + 1) / len(study_cards)
                st.progress(progress, text=f"Card {st.session_state.study_index + 1} of {len(study_cards)}")

                # Front of card
                st.markdown(f"""
                <div style="
                    border: 2px solid #ddd;
                    border-radius: 10px;
                    padding: 30px;
                    margin: 20px 0;
                    background-color: #f9f9f9;
                    color: black;
                    text-align: center;
                    min-height: 150px;
                ">
                    <h3>{current_card['front']}</h3>
                </div>
                """, unsafe_allow_html=True)

                # Show answer + grading buttons
                if st.session_state.show_answer:
                    st.markdown(f"""
                    <div style="border: 2px solid #4CAF50; border-radius: 10px; padding: 20px; margin: 20px 0; 
                               background-color: #e8f5e8; color: black; text-align: center;">
                        <h4>Answer:</h4>
                        <p>{current_card['back']}</p>
                    </div>
                    """, unsafe_allow_html=True)

                    st.markdown("### How well did you know this?")
                    c1, c2, c3 = st.columns(3)

                    with c1:
                        if st.button("❌ Needs work", use_container_width=True):
                            user_data.save_current_user(st.session_state)
                            next_flashcard(study_cards, correct=False)

                    with c2:
                        if st.button("🤔 Almost", use_container_width=True):
                            user_data.save_current_user(st.session_state)
                            next_flashcard(study_cards, correct=True)

                    with c3:
                        if st.button("✅ Mastered", use_container_width=True):
                            user_data.save_current_user(st.session_state)
                            next_flashcard(study_cards, correct=True)
                else:
                    if st.button("🔍 Show Answer", use_container_width=True):
                        st.session_state.show_answer = True
                        st.rerun()

                # Session accuracy m
                if st.session_state.cards_studied > 0:
                    accuracy = (st.session_state.cards_correct / st.session_state.cards_studied) * 100
                    st.metric("Session Accuracy", f"{accuracy:.1f}%")

    # Create Tab
    with tab2:
        st.subheader("➕ Create Flashcards")

        method = st.radio(
            "Creation method:",
            ["📝 From Text", "📂 Upload File", "✋ Manual Entry", "📚 From Notes"],
            horizontal=True
        )

        # From Text
        if method == "📝 From Text":
            content = st.text_area("Paste content:", placeholder="Enter study material...", height=150)
            c1, c2, c3 = st.columns(3)
            with c1:
                num_cards = st.slider("Number of cards:", 3, 20, 8)
            with c2:
                difficulty = st.selectbox("Difficulty:", ["Easy", "Medium", "Hard"])
            with c3:
                category = st.text_input("Category:", value="General")

            if st.button("🚀 Generate Flashcards", type="primary"):
                if content.strip():
                    with st.spinner("Creating flashcards..."):
                        try:
                            flashcards = generators['flashcards'].generate_flashcards(
                                content, num_cards=num_cards, difficulty=difficulty
                            )
                            for card in flashcards:
                                card['category'] = category

                            st.session_state.flashcards.extend(flashcards)
                            auto_save()
                            st.success(f"✅ Generated {len(flashcards)} flashcards!")

                            # Log activity
                            session = {
                                'timestamp': datetime.now().isoformat(),
                                'activity_type': 'flashcards_created',
                                'subject': category,
                                'flashcards_created': len(flashcards)
                            }
                            st.session_state.study_sessions.append(session)
                            auto_save()

                            # Preview some cards
                            st.markdown("### Preview:")
                            for i, card in enumerate(flashcards[:3], 1):
                                with st.expander(f"Card {i}"):
                                    st.write(f"**Front:** {card['front']}")
                                    st.write(f"**Back:** {card['back']}")
                            if len(flashcards) > 3:
                                st.info(f"+ {len(flashcards) - 3} more cards created!")
                                ser_data.save_current_user(st.session_state)
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
                else:
                    st.warning("Please enter content.")

        # Manual Entry
        elif method == "✋ Manual Entry":
            with st.form("manual_flashcard"):
                front = st.text_area("Front (Question):", height=100)
                back = st.text_area("Back (Answer):", height=100)
                category = st.text_input("Category:", value="General")

                if st.form_submit_button("➕ Add Flashcard"):
                    user_data.save_current_user(st.session_state)
                    if front.strip() and back.strip():
                        new_card = {
                            'front': front,
                            'back': back,
                            'category': category,
                            'created': datetime.now().isoformat()
                        }
                        st.session_state.flashcards.append(new_card)
                        auto_save()
                        ser_data.save_current_user(st.session_state)
                        st.success("✅ Flashcard added!")
                    else:
                        st.warning("Please fill in both sides.")

        # Upload File
        elif method == "📂 Upload File":
            uploaded_file = st.file_uploader("Choose a file", type=["txt", "pdf", "docx"])
            content = ""
            if uploaded_file is not None:
                if uploaded_file.type == "text/plain":
                    content = uploaded_file.read().decode("utf-8")
                elif uploaded_file.type == "application/pdf":
                    pdf = PdfReader(uploaded_file)
                    content = "".join(page.extract_text() or "" for page in pdf.pages)
                elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    doc = docx.Document(uploaded_file)
                    content = "\n".join([p.text for p in doc.paragraphs])
                st.text_area("Preview:", value=(content[:200] + "...") if content else "", height=100, disabled=True)

            c1, c2, c3 = st.columns(3)
            with c1:
                num_cards = st.slider("Number of cards:", 3, 20, 8)
            with c2:
                difficulty = st.selectbox("Difficulty:", ["Easy", "Medium", "Hard"])
            with c3:
                category = st.text_input("Category:", value="General")

            if st.button("🚀 Generate Flashcards", type="primary"):
                user_data.save_current_user(st.session_state)
                if content.strip():
                    with st.spinner("Creating flashcards..."):
                        try:
                            flashcards = generators['flashcards'].generate_flashcards(
                                content, num_cards=num_cards, difficulty=difficulty
                            )
                            for card in flashcards:
                                card['category'] = category

                            st.session_state.flashcards.extend(flashcards)
                            auto_save()
                            user_data.save_current_user(st.session_state)
                            st.success(f"✅ Generated {len(flashcards)} flashcards!")

                            # Log activity
                            session = {
                                'timestamp': datetime.now().isoformat(),
                                'activity_type': 'flashcards_created',
                                'subject': category,
                                'flashcards_created': len(flashcards)
                            }
                            st.session_state.study_sessions.append(session)
                            auto_save()

                            # Preview some cards
                            st.markdown("### Preview:")
                            for i, card in enumerate(flashcards[:3], 1):
                                with st.expander(f"Card {i}"):
                                    st.write(f"**Front:** {card['front']}")
                                    st.write(f"**Back:** {card['back']}")
                            if len(flashcards) > 3:
                                st.info(f"+ {len(flashcards) - 3} more cards created!")
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
                else:
                    st.warning("Please enter content.")

        # From Notes
        elif method == "📚 From Notes":
            if st.session_state.notes:
                note_titles = [n['title'] for n in st.session_state.notes]
                selected_note = st.selectbox("Select note:", note_titles)
                note_obj = next((n for n in st.session_state.notes if n['title'] == selected_note), None)
                if note_obj:
                    content = note_obj['content']
                    st.text_area("Preview:", value=content[:200] + "...", height=100, disabled=True)

                    c1, c2, c3 = st.columns(3)
                    with c1:
                        num_cards = st.slider("Number of cards:", 3, 20, 8)
                    with c2:
                        difficulty = st.selectbox("Difficulty:", ["Easy", "Medium", "Hard"])
                    with c3:
                        category = st.text_input("Category:", value="General")

                    if st.button("🚀 Generate Flashcards from Note", type="primary"):
                        user_data.save_current_user(st.session_state)
                        if content.strip():
                            with st.spinner("Creating flashcards..."):
                                try:
                                    flashcards = generators['flashcards'].generate_flashcards(
                                        content, num_cards=num_cards, difficulty=difficulty
                                    )
                                    for card in flashcards:
                                        card['category'] = category

                                    st.session_state.flashcards.extend(flashcards)
                                    auto_save()
                                    user_data.save_current_user(st.session_state)
                                    st.success(f"✅ Generated {len(flashcards)} flashcards from note!")

                                    # Log activity
                                    session = {
                                        'timestamp': datetime.now().isoformat(),
                                        'activity_type': 'flashcards_created',
                                        'subject': category,
                                        'flashcards_created': len(flashcards)
                                    }
                                    st.session_state.study_sessions.append(session)
                                    auto_save()

                                    # Preview
                                    st.markdown("### Preview:")
                                    for i, card in enumerate(flashcards[:3], 1):
                                        with st.expander(f"Card {i}"):
                                            st.write(f"**Front:** {card['front']}")
                                            st.write(f"**Back:** {card['back']}")
                                    if len(flashcards) > 3:
                                        st.info(f"+ {len(flashcards) - 3} more cards created!")
                                except Exception as e:
                                    st.error(f"Error: {str(e)}")
                        else:
                            st.warning("Note is empty.")
            else:
                st.info("No notes available. Create some first!")

    # Manage Tab
    with tab3:
        st.subheader("📂 Manage Flashcards")

        if st.session_state.flashcards:
            st.write(f"**Total flashcards:** {len(st.session_state.flashcards)}")

            # Export / Clear controls
            c1, c2 = st.columns(2)
            with c1:
                if st.button("📥 Export All"):
                    data = generators['flashcards'].save_flashcards_file(
                        st.session_state.flashcards, "export"
                    )
                    st.download_button(
                        "Download",
                        data=data,
                        file_name=f"flashcards_{datetime.now().strftime('%Y%m%d')}.json",
                        mime="application/json"
                    )
            with c2:
                if st.button("🗑️ Clear All"):
                    user_data.save_current_user(st.session_state)
                    st.session_state.flashcards = []
                    auto_save()
                    st.success("✅ All flashcards deleted!")
                    st.rerun()

            # Filter and display
            categories = list(set([card.get('category', 'General') for card in st.session_state.flashcards]))
            filter_cat = st.selectbox("Filter:", ["All"] + categories)

            filtered = st.session_state.flashcards
            if filter_cat != "All":
                filtered = [c for c in filtered if c.get('category', 'General') == filter_cat]

            for i, card in enumerate(filtered):
                with st.expander(f"🎴 {card['front'][:50]}..."):
                    st.write(f"**Front:** {card['front']}")
                    st.write(f"**Back:** {card['back']}")
                    st.write(f"**Category:** {card.get('category', 'General')}")

                    if st.button("🗑️ Delete", key=f"del_{i}"):
                        user_data.save_current_user(st.session_state)
                        st.session_state.flashcards.remove(card)
                        auto_save()
                        st.rerun()
        else:
            st.info("No flashcards yet. Create some first!")


# ============================
# Quizzes Page
# ============================

elif st.session_state.page == "🧠 Quizzes":
    st.title("🧠 Interactive Quiz System")

    tab1, tab2 = st.tabs(["📝 Take Quiz", "📊 History"])

    # Take Quiz
    with tab1:
        # If a quiz is active, display it via AdvancedQuizSystem
        if st.session_state.get('quiz_active', False):
            advanced_quiz.display_quiz_interface(st.session_state.get('current_quiz'))
        else:
            st.subheader("📝 Create New Quiz")

            content = ""  # Will hold the text used to generate the quiz

            c1, c2 = st.columns(2)
            with c1:
                source = st.radio("Quiz source:", ["📚 My Notes", "📝 New Content", "📂 Upload file"])
                question_type = st.selectbox(
                    "Question Type:",
                    options=[
                        "Multiple Choice Only",
                        "True/False Only",
                        "Short Answer Only",
                        "Mixed Questions"
                    ],
                    index=3
                )
            with c2:
                num_questions = st.slider("Questions:", 3, 15, 8)
                difficulty = st.selectbox("Difficulty:", ["Easy", "Medium", "Hard"])

            # Choose content source
            if source == "📚 My Notes":
                if st.session_state.notes:
                    note_titles = [n['title'] for n in st.session_state.notes]
                    selected_note = st.selectbox("Select note:", note_titles)
                    note_obj = next((n for n in st.session_state.notes if n['title'] == selected_note), None)
                    if note_obj:
                        content = note_obj['content']
                        st.text_area("Preview:", value=content[:200] + "...", height=100, disabled=True)
                else:
                    st.info("No notes available. Create some first!")

            elif source == "📝 New Content":
                content = st.text_area("Enter content:", height=200, placeholder="Paste study material...")

            elif source == "📂 Upload file":
                uploaded_file = st.file_uploader("Choose a file", type=["txt", "pdf", "docx"])
                if uploaded_file is not None:
                    if uploaded_file.type == "text/plain":
                        content = uploaded_file.read().decode("utf-8")
                    elif uploaded_file.type == "application/pdf":
                        pdf = PdfReader(uploaded_file)
                        content = "".join(page.extract_text() or "" for page in pdf.pages)
                    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        doc = docx.Document(uploaded_file)
                        content = "\n".join([p.text for p in doc.paragraphs])
                    st.text_area("Preview:", value=content[:200] + "...", height=100, disabled=True)

            # Start quiz
            if st.button("🚀 Create & Start Quiz", type="primary", use_container_width=True):
                user_data.save_current_user(st.session_state)
                if content.strip():
                    with st.spinner("Creating quiz..."):
                        try:
                            quiz_data = advanced_quiz.create_quiz_from_content(
                                content,
                                num_questions=num_questions,
                                difficulty=difficulty,
                                question_type=question_type
                            )
                            if quiz_data and quiz_data.get('questions'):
                                st.session_state.current_quiz = quiz_data
                                st.session_state.quiz_active = True
                                st.session_state.quiz_question_index = 0
                                st.session_state.quiz_answers = {}
                                st.success("✅ Quiz created! Starting now...")
                                st.rerun()
                            else:
                                st.error("Failed to create quiz. Please try again.")
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
                else:
                    st.warning("Please provide content for the quiz.")

    # History
    with tab2:
        st.subheader("📊 Quiz History")

        quiz_sessions = [s for s in st.session_state.study_sessions if s.get('activity_type') == 'quiz']

        if quiz_sessions:
            c1, c2, c3 = st.columns(3)
            with c1:
                st.metric("Total Quizzes", len(quiz_sessions))
            with c2:
                avg_score = sum(s.get('score', 0) for s in quiz_sessions) / len(quiz_sessions)
                st.metric("Average Score", f"{avg_score:.1f}%")
            with c3:
                best_score = max(s.get('score', 0) for s in quiz_sessions)
                st.metric("Best Score", f"{best_score:.1f}%")

            st.subheader("Recent Results")
            for i, session in enumerate(reversed(quiz_sessions[-10:])):
                timestamp = datetime.fromisoformat(session['timestamp']).strftime("%Y-%m-%d %H:%M")
                score = session.get('score', 0)
                correct = session.get('correct_answers', 0)
                total = session.get('total_questions', 0)
                color = "🟢" if score >= 80 else "🟡" if score >= 60 else "🔴"

                with st.expander(f"{color} {timestamp} - {score:.1f}% ({correct}/{total})"):
                    st.write(f"**Score:** {score:.1f}%")
                    st.write(f"**Difficulty:** {session.get('difficulty', 'N/A')}")

                    # Retake
                    if st.button("🔄 Retake This Quiz", key=f"retake_{i}"):
                        user_data.save_current_user(st.session_state)
                        st.session_state.retake_quiz_content = session.get('original_content', '')
                        st.session_state.retake_quiz_config = {
                            'num_questions': session.get('total_questions', 10),
                            'difficulty': session.get('difficulty', 'Medium')
                        }
                        st.session_state.page = "🧠 Quizzes"
                        st.session_state.quiz_active = False
                        st.rerun()
        else:
            st.info("No quiz history yet. Take your first quiz!")


# ============================
# Progress Page
# ============================

elif st.session_state.page == "📊 Progress":
    st.title("📊 Learning Analytics")

    if not st.session_state.study_sessions:
        st.info("No study data yet. Start using the platform to see your progress!")
    else:
        # Summary metrics
        c1, c2, c3, c4 = st.columns(4)

        total_sessions = len(st.session_state.study_sessions)
        quiz_sessions = [s for s in st.session_state.study_sessions if s.get('activity_type') == 'quiz']
        flashcard_sessions = [s for s in st.session_state.study_sessions if s.get('activity_type') in ['flashcards', 'flashcards_created']]

        with c1:
            st.metric("Total Sessions", total_sessions)
        with c2:
            st.metric("Quizzes Taken", len(quiz_sessions))
        with c3:
            if quiz_sessions:
                avg_quiz_score = sum(s.get('score', 0) for s in quiz_sessions) / len(quiz_sessions)
                st.metric("Avg Quiz Score", f"{avg_quiz_score:.1f}%")
            else:
                st.metric("Avg Quiz Score", "N/A")
        with c4:
            st.metric("Flashcard Sessions", len(flashcard_sessions))

        # Recent activity bar chart (last 7 days)
        st.subheader("📈 Recent Activity")

        daily_activity = defaultdict(int)
        for session in st.session_state.study_sessions:
            try:
                date = datetime.fromisoformat(session['timestamp']).strftime('%Y-%m-%d')
                daily_activity[date] += 1
            except Exception:
                # Skip malformed timestamps
                continue

        if daily_activity:
            dates = list(daily_activity.keys())[-7:]
            counts = [daily_activity[date] for date in dates]

            import matplotlib.pyplot as plt  # Imported here to keep top imports minimal
            fig, ax = plt.subplots(figsize=(10, 4))
            ax.bar(dates, counts)
            ax.set_title('Study Sessions (Last 7 Days)')
            ax.set_ylabel('Sessions')
            plt.xticks(rotation=45)
            st.pyplot(fig)

# ============================
# Calendar Page
# ============================

elif st.session_state.page == "📅 Calendar":
    st.title("📅 Calendar & Events")

    # Initialize calendar sessions
    if "events" not in st.session_state:
        st.session_state.events = []
    if "calendar_year" not in st.session_state:
        st.session_state.calendar_year = datetime.now().year
    if "calendar_month" not in st.session_state:
        st.session_state.calendar_month = datetime.now().month
    if "selected_date" not in st.session_state:
        st.session_state.selected_date = None

    # Month navigation helper
    def change_month(delta):
        new_month = st.session_state.calendar_month + delta
        if new_month > 12:
            st.session_state.calendar_month = 1
            st.session_state.calendar_year += 1
        elif new_month < 1:
            st.session_state.calendar_month = 12
            st.session_state.calendar_year -= 1
        else:
            st.session_state.calendar_month = new_month

    # Add Event form
    st.subheader("➕ Add Event")
    with st.form("add_event_form"):
        user_data.save_current_user(st.session_state)
        name = st.text_input("Event Title:", placeholder="e.g., Math Test, History Project, Concert")
        date = st.date_input("Date:")
        notes = st.text_area("Details (optional):", placeholder="Extra info...")
        color = st.color_picker("Pick a color:", "#4CAF50")  # default green

        submitted = st.form_submit_button("➕ Add")
        if submitted:
            if name.strip():
                new_event = {
                    "name": name,
                    "date": date.isoformat(),
                    "notes": notes,
                    "color": color,
                    "created": datetime.now().isoformat()
                }
                st.session_state.events.append(new_event)
                user_data.save_current_user(st.session_state)
                auto_save()
                st.success(f"✅ Added event - {name}")
            else:
                st.warning("Please enter a title.")

    st.divider()

    # Month Navigation Controls
    c1, c2, c3 = st.columns([1, 3, 1])
    with c1:
        st.button("‹", key="prev_month", on_click=change_month, args=(-1,))
    with c2:
        month_name = datetime(
            st.session_state.calendar_year,
            st.session_state.calendar_month,
            1
        ).strftime("%B %Y")
        st.markdown(f"<h2 style='text-align:center;margin:0'>{month_name}</h2>", unsafe_allow_html=True)
    with c3:
        st.button("›", key="next_month", on_click=change_month, args=(1,))

    st.divider()

    # Build the month grid
    cal = calendar.Calendar(firstweekday=0)  # Monday start (0=Monday)
    month_days = list(cal.itermonthdates(st.session_state.calendar_year, st.session_state.calendar_month))

    # Inline CSS + HTML grid layout
    html = """
    <style>
        .calendar-day {
            padding: 10px;
            background: #f9f9f9;
            border-radius: 8px;
            border: 1px solid #ddd;
            min-height: 80px;
            text-align: left;
            position: relative;
            transition: all 0.2s ease;
        }
        .calendar-day:hover {
            background: #f0f0f0;
            transform: translateY(-2px);
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }
        .day-number {
            position: absolute;
            top: 5px;
            right: 8px;
            font-weight: 600;
            font-size: 14px;
            color: #333;
        }
        .today-highlight {
            border: 2px solid #2196F3 !important;
            background: #e3f2fd !important;
        }
        .event-item {
            margin: 3px 0;
            border-radius: 4px;
            font-size: 11px;
            padding: 3px 5px;
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
            box-shadow: 0 1px 2px rgba(0,0,0,0.1);
        }
        .event-notes {
            font-size: 10px;
            color: #555;
            margin-top: 2px;
            white-space: normal;
        }
        .more-events {
            font-size: 10px;
            color: #666;
            margin-top: 3px;
            font-style: italic;
        }
    </style>
    <div style='display:grid; grid-template-columns: repeat(7, 1fr); gap:8px; font-family:sans-serif; text-align:center; width:100%;'>
    """

    # Weekday headers
    for wd in ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]:
        html += f"<div style='font-weight:bold; padding:10px; background:#f0f0f0; border-radius:4px; color: black;'>{wd}</div>"

    today = datetime.now().date()

    # Render each day cell with events
    for day in month_days:
        if day.month == st.session_state.calendar_month:
            events_today = [e for e in st.session_state.events if datetime.fromisoformat(e["date"]).date() == day]

            today_class = "today-highlight" if day == today else ""

            html += f"<div class='calendar-day {today_class}'>"
            html += f"<div class='day-number'>{day.day}</div>"

            if events_today:
                html += "<div style='margin-top:20px; max-height:60px; overflow-y:auto; padding-right:3px;'>"
                for e in events_today[:4]:
                    html += f"<div class='event-item' style='background:{e['color']}; color:white;' title='{e['name']}'>"
                    html += f"{e['name']}"
                    if e.get('notes'):
                        html += f"<div class='event-notes'>{e['notes']}</div>"
                    html += "</div>"
                if len(events_today) > 4:
                    html += f"<div class='more-events'>+{len(events_today)-4} more</div>"
                html += "</div>"

            html += "</div>"
        else:
            # Blank cell for days outside current month grid
            html += "<div style='padding:10px; color:#ccc; min-height:80px;'></div>"

    html += "</div>"

    # Display Calendar grid
    st.markdown(html, unsafe_allow_html=True)

    # Today's reminders
    st.divider()
    st.subheader("📅 Today's Reminders")
    today_events = [e for e in st.session_state.events if datetime.fromisoformat(e["date"]).date() == today]
    if today_events:
        for e in today_events:
            with st.container(border=True):
                st.markdown(f"<span style='color:{e['color']};font-size:20px'>●</span> **{e['name']}**", unsafe_allow_html=True)
                if e.get('notes'):
                    st.caption(f"📝 {e['notes']}")
    else:
        st.info("No events today.")

    # Upcoming reminders
    st.divider()
    st.subheader("⏰ Upcoming")
    upcoming = [e for e in st.session_state.events if datetime.fromisoformat(e["date"]).date() > today]
    upcoming = sorted(upcoming, key=lambda x: x["date"])[:5]
    if upcoming:
        for e in upcoming:
            with st.container(border=True):
                date_obj = datetime.fromisoformat(e["date"]).strftime("%a, %b %d")
                st.markdown(f"<span style='color:{e['color']};font-size:20px'>●</span> **{date_obj}** — {e['name']}", unsafe_allow_html=True)
                if e.get('notes'):
                    st.caption(f"📝 {e['notes']}")
    else:
        st.info("No upcoming events!")

    # Delete event UI
    st.divider()
    st.subheader("🗑️ Delete Event")

    if st.session_state.events:
        # Dropdown formatted as "YYYY/MM/DD — Event Name"
        event_options = [
            f"{datetime.fromisoformat(e['date']).strftime('%Y/%m/%d')} — {e['name']}"
            for e in st.session_state.events
        ]
        event_to_delete = st.selectbox("Select an event to delete:", options=event_options)

        if st.button("❌ Delete Selected Event"):
            user_data.save_current_user(st.session_state)
            for e in list(st.session_state.events):
                label = f"{datetime.fromisoformat(e['date']).strftime('%Y/%m/%d')} — {e['name']}"
                if label == event_to_delete:
                    st.session_state.events.remove(e)
                    auto_save()
                    st.success(f"Deleted event: {label}")
                    st.rerun()
                    user_data.save_current_user(st.session_state)
    else:
        st.info("No events to delete.")
elif st.session_state.page == "📝 Autograder":
    from autograder import AutoGrader
    st.title("📝 AI Autograder")

    text_input = st.text_area("✍️ Paste your essay, story, or text:", height=250, placeholder="Write or paste your text here...")

    col1, col2 = st.columns(2)
    with col1:
        text_type = st.selectbox("Text type:", ["Essay", "Story", "Article", "Other"])
    with col2:
        extra_notes = st.text_input("Extra notes (optional)", placeholder="e.g., Focus on creativity, academic tone...")

    if st.button("🚀 Grade Now", type="primary", use_container_width=True):
        if not text_input.strip():
            st.warning("Please enter some text to grade.")
        else:
            grader = AutoGrader()
            with st.spinner("Grading with AI..."):
                result = grader.grade_text(text_input, text_type, extra_notes)

            # --- Stylish Results ---
            st.subheader(f"📊 Score: {result.get('score', 0)}/10")
            st.progress(int(result.get("score", 0)) / 10)

            st.markdown("### ✅ Strengths")
            for s in result.get("strengths", []):
                st.markdown(f"- {s}")

            st.markdown("### ⚠️ Weaknesses")
            for w in result.get("weaknesses", []):
                st.markdown(f"- {w}")

            st.markdown("### 💡 Suggestions to Improve")
            for sug in result.get("suggestions", []):
                st.markdown(f"- {sug}")

            st.markdown("### 📝 Detailed Feedback")
            st.info(result.get("detailed_feedback", "No feedback provided."))
# ============================
# Settings Page
# ============================

elif st.session_state.page == "⚙️ Settings":
    st.title("⚙️ User Settings")

    # --- Change Password Section ---
    st.subheader("🔑 Change Your Password")
    with st.form("change_password_form"):
        current_password = st.text_input("Current Password", type="password")
        new_password = st.text_input("New Password", type="password")
        confirm_new_password = st.text_input("Confirm New Password", type="password")
        
        submitted = st.form_submit_button("Change Password")
        
        if submitted:
            if not all([current_password, new_password, confirm_new_password]):
                st.warning("Please fill out all password fields.")
            elif new_password != confirm_new_password:
                st.error("New passwords do not match.")
            else:
                # Authenticate user with their current password first
                is_valid, _ = user_data.authenticate(st.session_state['username'], current_password)
                if is_valid:
                    # If authentication succeeds, proceed with the password update
                    success, msg = user_data.change_password(st.session_state['username'], new_password)
                    if success:
                        st.success("Your password has been changed successfully!")
                    else:
                        st.error(f"Failed to change password: {msg}")
                else:
                    st.error("The current password you entered is incorrect.")

    st.divider()

    # --- Delete Account Section ---
    st.subheader("🗑️ Danger Zone")
    with st.expander("Delete Account"):
        st.warning("This action is irreversible. All your notes, flashcards, and progress will be permanently deleted.")
        
        confirmation_text = st.text_input(
            "To confirm, please type `DELETE` in the box below:"
        )
        
        if st.button("Permanently Delete My Account", type="primary"):
            if confirmation_text == "DELETE":
                success, msg = user_data.delete_account(st.session_state['username'])
                if success:
                    st.success("Your account has been deleted. You have been logged out.")
                    # Clear session and log the user out
                    st.session_state.clear()
                    st.session_state["logged_in"] = False
                    st.session_state["page"] = "🏠 Home"
                    st.rerun()
                else:
                    st.error(f"An error occurred while deleting your account: {msg}")
            else:
                st.error("Confirmation text did not match. Account deletion cancelled.")




# ============================
# Autosave hook 
# ============================

# Every 5 sessions, trigger an autosave (kept same behavior)
if len(st.session_state.study_sessions) % 5 == 0:
    auto_save()
