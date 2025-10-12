#BACKUP ONLY
#Author: Victor
# Page name: app.py
# Page purpose: Main app file for the AI Study Platform using Streamlit
# Date of creation: 2025-10-10

#imports
import io
import re
import random                       # For choosing random numbers
import calendar                     # For building the calendar view
import hashlib                      # For hashing passwords
from datetime import datetime       # For dates
from datetime import timedelta      # For duration 
from collections import defaultdict
import streamlit as st            
import matplotlib.pyplot as plt  
from PyPDF2 import PdfReader        # PDF text extraction
import docx                         # Docx text extraction

from supabase import create_client  # Supabase client

# Local Imports
import user_data
from note_generator import NoteGenerator
from flashcard_generator import FlashcardGenerator
from quiz_generator import QuizGenerator
from progress_tracker import ProgressTracker
from pdf_report_generator import PDFReportGenerator
from data_persistence import DataPersistence
from advanced_quiz_system import AdvancedQuizSystem
from data_import_export import DataImportExport
from utils import sanitize_filename
from autograder import AutoGrader



#-----------------------
# Configuration & Setup
#-----------------------

# App base configuration
st.set_page_config(
    page_title="AI Study Platform",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Secrets & Keys
SUPABASE_URL = st.secrets["SUPABASE_URL"]
SUPABASE_KEY = st.secrets["SUPABASE_KEY"]
ADMIN_KEY = st.secrets["ADMIN_KEY"]  # Used for hidden admin mode toggle

# Supabase Client
supabase = create_client(SUPABASE_URL, SUPABASE_KEY)



# Functions

# Hash a raw password into SHA-256
def hash_password(password: str) -> str:
    # Returns a hex string hash for secure storage
    return hashlib.sha256(password.encode()).hexdigest()

# Delete any user's account, admin only.
def admin_delete_account(target_username: str):
    # Attempts to delete a user row from Supabase 'users' table
    try:
        supabase.from_("users").delete().eq("username", target_username).execute()
        st.success(f"✅ Account '{target_username}' has been deleted!")
    except Exception as e:
        st.error(f"Error: {e}")


# Reset any user's password (Admin-only)
def admin_reset_password(target_username: str, new_password: str):
    # Hash the new password and update the user's record
    try:
        hashed = hash_password(new_password)
        # FIXED: call lower() properly
        supabase.table("users").update({"password": hashed}).eq("username", target_username.lower()).execute()
        st.success(f"✅ Password for '{target_username}' has been reset!")
    except Exception as e:
        st.error(f"Error: {e}")

# Advance the flashcard study session to the next card
def next_flashcard(study_cards, correct=False):
    # Increment counters for studied and correct answers
    st.session_state.cards_studied += 1
    if correct:
        st.session_state.cards_correct += 1

    # Hide answer for the next card
    st.session_state.study_index += 1
    st.session_state.show_answer = False

    # If we reached the end, finalize the session
    if st.session_state.study_index >= len(study_cards):
        # Compute accuracy for this study session
        accuracy = (st.session_state.cards_correct / st.session_state.cards_studied) * 100

        # Save a "flashcards" study session record
        session = {
            'timestamp': datetime.now().isoformat(),
            'activity_type': 'flashcards',
            'subject': study_cards[0].get('category', 'General') if study_cards else 'General',
            'flashcards_studied': st.session_state.cards_studied,
            'correct_answers': st.session_state.cards_correct,
            'accuracy': accuracy
        }
        st.session_state.study_sessions.append(session)
        auto_save()

        st.success(f"Session complete! Accuracy: {accuracy:.1f}%")

        # Clean up study session keys
        for key in ['study_index', 'show_answer', 'cards_studied', 'cards_correct']:
            if key in st.session_state:
                del st.session_state[key]

    # Refresh
    st.rerun()


# Autosave
def auto_save():
    # Only autosave if not rerun (avoid on every rerun for speed)
    try:
        persistence.auto_save_data()
    except Exception:
        pass

# Resources

# Cache only resource creation
@st.cache_resource
def get_generators():
    return {
        'notes': NoteGenerator(),
        'flashcards': FlashcardGenerator(),
        'quiz': QuizGenerator(),
        'progress': ProgressTracker(),
        'pdf': PDFReportGenerator()
    }

@st.cache_resource
def get_persistence():
    return DataPersistence()

# Names the functions
generators = get_generators()
persistence = get_persistence()
data_io = DataImportExport(persistence)
advanced_quiz = AdvancedQuizSystem(generators['quiz'])

# Session State Initialization

def init_session_state():
    # Provide defaults for keys we rely upon throughout the app
    defaults = {
        'notes': [],
        'flashcards': [],
        'study_sessions': [],
        'current_note': "",
        'note_title': "",
        'note_category': "General",
        'page': 'Home',
        'logged_in': False,
        'username': ""
    }
    for key, default_value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = default_value

init_session_state()

persistence.load_all_data()

# ------------------------------
# Ensure page state always exists
# -----------------------------
if "page" not in st.session_state:
    st.session_state.page = "🏠 Home"

# ----------------------------
# Sidebar
# ----------------------------
with st.sidebar:
    st.title("🎓 Study Platform")

    # ===========================
    # Logged Out → Login / Sign Up
    # ===========================
    if not st.session_state.get("logged_in", False):
        st.subheader("🔑 Account Access")

        mode = st.radio("Choose mode", ["Login", "Sign Up"], horizontal=True, key="auth_mode")

        if mode == "Sign Up": # Sign Up
            su = st.text_input("Username", max_chars=15, key="su_user")
            sp = st.text_input("Password", type="password", key="su_pass")
            confirm = st.text_input("Confirm Password", type="password", key="su_confirm")

            if st.button("Create account", use_container_width=True):
                if not su or not sp:
                    st.warning("Enter username and password.")
                elif len(su) > 15 or len(su) < 4:
                    st.warning("Username must be between 4 and 15 characters.")
                elif ' ' in su:
                    st.error("Username cannot contain spaces")
                elif not su.isalnum():
                    st.error("Username can only contain letters and numbers")
                elif sp != confirm:
                    st.error("Passwords do not match.")
                else:
                    ok, msg = user_data.register_user(su, sp)
                    if ok:
                        st.success("Account created. You can now log in.")
                    else:
                        st.error(msg)


        else:  # Login
            lu = st.text_input("Username", key="li_user")
            lp = st.text_input("Password", type="password", key="li_pass")

            if st.button("Login", use_container_width=True):
                if lu.strip() == "" and lp == st.secrets.get("ADMIN_KEY", ""):
                    st.session_state["admin_mode"] = True
                    st.session_state["logged_in"] = True
                    st.session_state["username"] = "Admin"
                    st.success("✅ Admin mode enabled")
                    st.rerun()
                else:
                    ok, msg = user_data.authenticate(lu, lp)
                    if ok:
                        st.session_state["logged_in"] = True
                        st.session_state["username"] = lu
                        loaded_ok, data = user_data.load_user_data(
                            lu, merge_local=True, local_state=st.session_state
                        )
                        if loaded_ok:
                            st.session_state.update(data)
                            user_data.save_current_user(st.session_state)
                        st.success(f"Welcome back, {lu}")
                        st.rerun()
                    else:
                        st.error(msg)

        # Force Home page when logged out
        st.session_state.page = "🏠 Home"

    # ===========================
    # Logged In → Sidebar Nav, Stats, Save
    # ===========================
    else:
        # Navigation at the top
        page = st.selectbox(
            "Navigate:",
            ["🏠 Home", "📝 Notes", "📚 Flashcards", "🧠 Quizzes",
             "📊 Progress", "📅 Calendar", "📝 Autograder", "⚙️ Settings"],
            key="navigation"
        )
        st.session_state.page = page
        # Welcome message
        st.subheader(f"👋 Welcome, {st.session_state['username']}")
        if st.session_state.get("admin_mode"):
            st.caption("🛠 Admin Mode Enabled")

        # Save button
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button("Save", use_container_width=True):
                ok, msg = user_data.save_current_user(st.session_state)
                if ok:
                    st.success("Saved.")
                else:
                    st.error(msg)
        # Quick Stats
        st.subheader("📈 Quick Stats")
        c1, c2 = st.columns(2)
        with c1:
            st.metric("📝 Notes", len(st.session_state.notes))
            st.metric("🎴 Flashcards", len(st.session_state.flashcards))
        with c2:
            quiz_sessions = [
                s for s in st.session_state.study_sessions
                if s.get('activity_type') == 'quiz'
            ]
            st.metric("🧠 Quizzes", len(quiz_sessions))
            st.metric("📚 Sessions", len(st.session_state.study_sessions))

        # Admin mode indicator
        if st.session_state.get("admin_mode"):
            st.info("🛠 Admin Mode Active")

        # not added currently data_io.render_sidebar_controls()

        # Logout
        if st.button("Logout", use_container_width=True):
            st.session_state.clear()
            st.session_state["logged_in"] = False
            st.session_state["page"] = "🏠 Home"
            st.success("Logged out.")
            st.rerun()


# ----------------------------
# Home Page When logged in
# ----------------------------
if st.session_state.page == "🏠 Home":
    # CSS styling
    st.markdown("""
    <style>
    .minimal-header {
        padding: 25px 0;
        margin-bottom: 30px;
        border-bottom: 1px solid #eaeaea;
    }
    .stat-card {
        background: white;
        padding: 20px;
        border-radius: 8px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.08);
        text-align: center;
        border: 1px solid #f0f0f0;
        transition: all 0.2s ease;
    }
    .stat-card:hover {
        box-shadow: 0 4px 12px rgba(0,0,0,0.1);
        transform: translateY(-2px);
    }
    .activity-item {
        background: white;
        padding: 16px;
        border-radius: 6px;
        margin-bottom: 10px;
        border-left: 3px solid #667eea;
        box-shadow: 0 1px 4px rgba(0,0,0,0.04);
    }
    .quick-action-btn {
        background: white;
        color: #333;
        border: 1px solid #e0e0e0;
        padding: 12px;
        border-radius: 8px;
        font-weight: 500;
        transition: all 0.2s ease;
        width: 100%;
    }
    .quick-action-btn:hover {
        border-color: #667eea;
        background: #f8f9ff;
    }
    .section-title {
        font-weight: 600;
        font-size: 1.1rem;
        margin-bottom: 15px;
        color: #333;
        padding-bottom: 8px;
        border-bottom: 1px solid #f0f0f0;
    }
    .admin-panel {
        background: #f8f9fa;
        border-radius: 8px;
        padding: 20px;
        margin-top: 30px;
        border: 1px solid #e9ecef;
    }
    </style>
    """, unsafe_allow_html=True)

    # Logged out so show app info
    if not st.session_state.get("logged_in", False):
        st.title("🎓 SmartStudy")
        st.markdown("---")
        st.subheader("Your AI-Powered Study Companion")
        
        col1, col2 = st.columns([2, 1])
        with col1:
            st.markdown("""
            ##### Transform Your Learning Experience
            
            SmartStudy helps you study smarter with AI-powered tools:
            
            - **📝 Smart Notes** - Generate and organize study materials
            - **🎴 Flashcards** - Create interactive study cards
            - **🧠 Quizzes** - Test your knowledge with adaptive quizzes
            - **📊 Analytics** - Track your progress and performance
            - **📅 Calendar** - Schedule study sessions and deadlines
            
            ##### Get Started:
            1. Create an account or log in
            2. Explore the study tools
            3. Track your learning journey
            """)
        
        with col2: # Image
            st.image("https://images.unsplash.com/photo-1501504905252-473c47e087f8?ixlib=rb-1.2.1&auto=format&fit=crop&w=500&q=80", 
                    caption="Study Smarter")

    # Logged in dashboard
    else:
        # Stylish minimal header
        st.markdown(f"<h1 style='color: white; margin-bottom: 8px;'>Welcome {st.session_state['username']}</h1>", unsafe_allow_html=True)
        st.markdown("<p style='color: #d1d5db; font-size: 1.1rem; margin-top: 0;'>Welcome back to your learning dashboard</p>", unsafe_allow_html=True)
        # Main content 
        col_left, col_right = st.columns([2, 1])
        
        with col_left:
            # Recent activity
            st.markdown("---")
            st.markdown('<h3 style="color: white; margin-bottom: 20px;">📅 Recent Activity</h3>', unsafe_allow_html=True)
            # Load recent 5 study sessions
            if st.session_state.get("study_sessions"):
                recent_sessions = sorted(
                    st.session_state.study_sessions,
                    key=lambda x: x.get('timestamp', ''), 
                    reverse=True
                )[:5]
                
                for i, session in enumerate(recent_sessions):
                    try:
                        timestamp = datetime.fromisoformat(session['timestamp']).strftime("%b %d %H:%M")
                        activity = session.get('activity_type', 'Unknown')
                        # Render based on activity type
                        if activity == 'quiz':
                            score = session.get('score', 0)
                            st.markdown(f"""
                            <div class="activity-item">
                                <div style="display: flex; justify-content: space-between; align-items: center;">
                                    <div style="font-weight: 500;">Quiz: {session.get('subject', 'General')}</div>
                                    <div style="color: { '#4caf50' if score >= 70 else '#ff9800' if score >= 50 else '#f44336' };">{score:.0f}%</div>
                                </div>
                                <div style="color: #888; font-size: 0.85rem; margin-top: 5px;">{timestamp}</div>
                            </div>
                            """, unsafe_allow_html=True)
                        
                        elif activity == 'flashcards': # Flashcards
                            studied = session.get('flashcards_studied', 0)
                            correct = session.get('correct_answers', 0)
                            accuracy = (correct / studied * 100) if studied > 0 else 0
                            st.markdown(f"""
                            <div class="activity-item">
                                <div style="display: flex; justify-content: space-between; align-items: center;">
                                    <div style="font-weight: 500;">Flashcards: {session.get('subject', 'General')}</div>
                                    <div>{studied} cards</div>
                                </div>
                                <div style="color: #888; font-size: 0.85rem; margin-top: 5px;">{timestamp} • {accuracy:.1f}% accuracy</div>
                            </div>
                            """, unsafe_allow_html=True)
                            
                    except:
                        continue
            else:
                st.info("No recent activity. Start studying to see your progress here.")
                
            # Quick note
            st.markdown("---")
            st.markdown('<h3 style="color: white; margin-bottom: 20px;">✏️ Quick Note</h3>', unsafe_allow_html=True)
            with st.form("quick_note_form"):
                quick_note = st.text_area("Jot something down:", placeholder="Type your quick note here...", height=100, 
                                         label_visibility="collapsed", key="quick_note_text")
                if st.form_submit_button("Save Note", use_container_width=True):
                    if quick_note.strip():
                        new_note = {
                            "title": f"Quick Note - {datetime.now().strftime('%H:%M')}",
                            "content": quick_note,
                            "category": "Quick Notes",
                            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        }
                        st.session_state.notes.append(new_note)
                        
                        # Delete the note you just saved
                        #st.session_state.notes.pop()
                        
                        auto_save()
                        st.rerun()
                        st.success("Quick note saved!")

        
        with col_right:
            # Upcoming events
            if "events" in st.session_state and st.session_state.events:
                st.markdown("---")
                st.markdown('<h3 style="color: white; margin-bottom: 20px;">📅 Upcoming Events</h3>', unsafe_allow_html=True)
                
                today = datetime.now().date()
                upcoming_events = []
                
                for event in st.session_state.events: # Parse and filter
                    try:
                        event_date = datetime.fromisoformat(event["date"]).date()
                        if event_date >= today:
                            upcoming_events.append({
                                "name": event["name"],
                                "date": event_date,
                                "color": event.get("color", "#667eea"),
                                "notes": event.get("notes", "")
                            })
                    except:
                        continue
                
                upcoming_events.sort(key=lambda x: x["date"])
                upcoming_events = upcoming_events[:3]
                
                for event in upcoming_events: # Render each event for home page
                    days_until = (event["date"] - today).days
                    day_text = "Today" if days_until == 0 else f"{days_until}d"
                    
                    st.markdown(f"""
                    <div style='background-color: white; padding: 12px; border-radius: 6px; margin-bottom: 10px; border-left: 3px solid {event["color"]}; box-shadow: 0 1px 4px rgba(0,0,0,0.04)'>
                        <div style='font-weight: 500; color:black;'>{event["name"]}</div>
                        <div style='color: #888; font-size: 0.85rem; margin-top: 5px;'>
                            {event["date"].strftime("%b %d")} • {day_text}
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
            
            # Study tip
            st.markdown("---")
            st.markdown('<h3 style="color: white; margin-bottom: 20px;">💡 Study Tip</h3>', unsafe_allow_html=True)
            study_tips = [
                "Focus on understanding concepts rather than memorizing facts.",
                "Take regular breaks to maintain focus and retention.",
                "Teach what you've learned to someone else to reinforce knowledge.",
                "Create connections between new information and what you already know.",
                "Practice retrieval by testing yourself regularly.",
                "Space out your study sessions over time for better long-term retention.",
                "Find a quiet, dedicated study space free from distractions."
            ]
            
            
            tip = random.choice(study_tips) # select a random tip
            st.info(f"{tip}")

        # Admin Controls (only if admin mode)
        if st.session_state.get("admin_mode"):
            st.markdown("---")
            st.header("🛠 Admin Controls")
            admin_col1, admin_col2 = st.columns(2)
            
            with admin_col1: # Reset password
                st.markdown("#### Reset Password")
                target_user = st.text_input("Username to Reset Password", key="admin_reset_user")
                new_pass = st.text_input("New Password", type="password", key="admin_new_pass")
                if st.button("Reset User Password", key="admin_reset_btn"):
                    if target_user and new_pass:
                        ok, msg = user_data.admin_reset_password(target_user, new_pass)
                        if ok:
                            st.success(f"Password for '{target_user}' reset successfully.")
                        else:
                            st.error(msg)
                    else:
                        st.warning("Enter both username and new password.")

            with admin_col2: # Delete account
                st.markdown("#### Delete Account")
                del_user = st.text_input("Username to Delete", key="admin_del_user")
                if st.button("Delete Account", type="secondary", key="admin_del_btn"):
                    if del_user:
                        ok, msg = user_data.admin_delete_account(del_user)
                        if ok:
                            st.success(f"User '{del_user}' deleted successfully.")
                        else:
                            st.error(msg)
                    else:
                        st.warning("Enter a username to delete.")
            
            st.markdown("</div>", unsafe_allow_html=True)


# ============================
# Notes Page
# ============================


elif st.session_state.page == "📝 Notes":
    st.title("📝 AI Note Generator & Class Notes")

    # Ensure keys used by this page exist
    for key in ["free_note_text", "free_note_title", "summarize_option", "free_note_cat", "topic_input", "topic_cat_input"]:
        if key not in st.session_state:
            if "text" in key or "title" in key or "input" in key:
                st.session_state[key] = ""
            elif "cat" in key:
                st.session_state[key] = "General"
            elif "summarize_option" in key:
                st.session_state[key] = True

    # Freeform Notes Mode
    st.subheader("💡 Freeform Notes Mode")

    note_title = st.text_input(
        "Note Name:",
        placeholder="Enter a title for your note...",
        key="free_note_title"
    )

    free_note = st.text_area(
        "Type your notes here:",
        value="",
        placeholder="Write your class notes here...",
        height=200,
        key="free_note_text"
    )

    free_category = st.text_input(
        "Category for these notes:",
        value=st.session_state.get("free_note_cat", "General"),
        key="free_note_cat"
    )

    summarize_option = st.checkbox(
        "🧠 Summarize with AI",
        value=st.session_state.get("summarize_option", True),
        key="summarize_option"
    )

    if st.button("💾 Save Notes", key="save_free_note"):
        user_data.save_current_user(st.session_state)
        if not free_note.strip():
            st.warning("Please enter some text to save.")
        else:
            notes_content = free_note
            if summarize_option:
                # AI summarise using NoteGenerator
                with st.spinner("Summarizing notes with AI..."):
                    try:
                        notes_content = generators['notes'].generate_notes(notes_content)
                    except Exception as e:
                        st.error(f"AI summarization failed: {e}")

            final_title = (note_title or "").strip() or "Untitled Note"

            # Append new note
            new_note = {
                "title": final_title,
                "content": notes_content,
                "category": free_category or "General",
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            st.session_state.notes.append(new_note)
            user_data.save_current_user(st.session_state)
            auto_save()
            st.success(f"✅ Note '{final_title}' saved!")
            st.rerun()

    st.markdown("---")

    # AI Note Generation from topic ora file
    st.subheader("🚀 Generate AI Notes from Topic or File")
    col1, col2 = st.columns([2, 1])
    with col1:
        topic = st.text_input("📖 Topic or subject:", placeholder="e.g., Photosynthesis, WWII, Calculus...", key="topic_input")
    with col2:
        category = st.text_input("Category:", value=st.session_state.topic_cat_input, key="topic_cat_input")

    uploaded_file = st.file_uploader(
        "📂 Upload a file (.txt, .md, .pdf, .docx) to generate notes:",
        type=['txt', 'md', 'pdf', 'docx'],
        key="file_upload"
    )

    if st.button("🚀 Generate Notes", key="generate_ai_notes"):
        user_data.save_current_user(st.session_state)
        content_to_process = ""
        note_name = ""

        # Extract content from uploaded file
        if uploaded_file:
            file_ext = uploaded_file.name.split('.')[-1].lower()
            note_name = uploaded_file.name.rsplit('.', 1)[0]
            try:
                if file_ext in ['txt', 'md']:
                    content_to_process = uploaded_file.read().decode("utf-8")
                elif file_ext == 'pdf':
                    pdf = PdfReader(io.BytesIO(uploaded_file.read()))
                    content_to_process = "\n".join([page.extract_text() or "" for page in pdf.pages])
                elif file_ext == 'docx':
                    doc = docx.Document(io.BytesIO(uploaded_file.read()))
                    content_to_process = "\n".join([para.text for para in doc.paragraphs])
                else:
                    st.error("Unsupported file type.")
            except Exception as e:
                st.error(f"Failed to read file: {e}")
        elif topic.strip():
            # Use typed topic directly
            content_to_process = topic
            note_name = topic
            

        # Generate notes through NoteGenerator
        if content_to_process:
            with st.spinner("Generating comprehensive notes..."):
                try:
                    notes_content = generators['notes'].generate_notes(content_to_process)
                    if notes_content:
                        new_note = {
                            "title": note_name or f"Note {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                            "content": notes_content,
                            "category": category or "General",
                            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        }
                        st.session_state.notes.append(new_note)
                        auto_save()
                        st.success(f"✅ Notes generated successfully for '{note_name}'!")
                        with st.expander("📖 Preview Generated Notes", expanded=True):
                            st.markdown(notes_content)
                            user_data.save_current_user(st.session_state)
                    else:
                        st.error("Failed to generate notes. Please try again.")
                except Exception as e:
                    st.error(f"Error: {e}")
        else:
            st.warning("Please enter a topic or upload a file.")
    # Existing Notes List
    if st.session_state.notes:
        st.divider()
        st.subheader("📚 Your Notes")

        categories = list(set([note.get('category', 'General') for note in st.session_state.notes]))
        filter_category = st.selectbox("Filter by category:", ["All"] + categories)

        filtered_notes = st.session_state.notes
        if filter_category != "All":
            filtered_notes = [n for n in filtered_notes if n.get('category') == filter_category]

        for i, note in enumerate(filtered_notes):
            with st.expander(f"📄 {note['title']} ({note.get('category', 'General')})"):
                st.write(f"**Created:** {note['timestamp']}")
                st.markdown(note['content'])

                col1, col2, col3, col4 = st.columns([1, 1, 1, 1])

                # Create flashcards from this note
                with col1:
                    if st.button("📚 Create Flashcards", key=f"flash_{i}"):
                        user_data.save_current_user(st.session_state)
                        with st.spinner("Creating flashcards..."):
                            try:
                                flashcards = generators['flashcards'].generate_flashcards(
                                    note['content'], num_cards=6, difficulty="Medium"
                                )
                                for card in flashcards:
                                    card['category'] = note.get('category', 'General')
                                st.session_state.flashcards.extend(flashcards)
                                auto_save()
                                st.success(f"✅ Created {len(flashcards)} flashcards!")

                                # Log activity
                                session = {
                                    "timestamp": datetime.now().isoformat(),
                                    "activity_type": "flashcards_created",
                                    "subject": note.get('category', 'General'),
                                    "flashcards_created": len(flashcards),
                                    "duration_minutes": 2
                                }
                                st.session_state.study_sessions.append(session)
                                auto_save()
                            except Exception as e:
                                st.error(f"Error: {e}")

                # Download note 
                with col2:
                    st.download_button(
                        "📥 Download",
                        data=note['content'],
                        file_name=f"{sanitize_filename(note['title'])}.txt",
                        mime="text/plain",
                        key=f"download_{i}"
                    )

                # Delete note
                with col3:
                    if st.button("🗑️ Delete", key=f"delete_{i}"):
                        st.session_state.notes.remove(note)
                        user_data.save_current_user(st.session_state)
                        auto_save()
                        st.rerun()

                # Rename note
                with col4:
                    new_name = st.text_input("Rename Note", value=note['title'], key=f"rename_{i}")
                    if st.button("✏️ Rename", key=f"rename_btn_{i}"):
                        user_data.save_current_user(st.session_state)
                        if new_name.strip():
                            note['title'] = new_name.strip()
                            auto_save()
                            st.success("✅ Note renamed!")
                            st.rerun()
                        else:
                            st.warning("Enter a valid name.")


# ============================
# Flashcards Page
# ============================

elif st.session_state.page == "📚 Flashcards":
    st.title("📚 Interactive Flashcards")

    tab1, tab2, tab3 = st.tabs(["📖 Study", "➕ Create", "📂 Manage"])

    # Study Tab
    with tab1:
        st.subheader("📖 Study Session")

        if not st.session_state.flashcards:
            st.info("No flashcards available. Create some first!")
            if st.button("🔄 Refresh"):
                user_data.save_current_user(st.session_state)
                st.rerun()
        else:
            # Filter by each category
            categories = list(set([card.get('category', 'General') for card in st.session_state.flashcards]))
            selected_category = st.selectbox("Study category:", ["All"] + categories)

            study_cards = st.session_state.flashcards
            if selected_category != "All":
                study_cards = [card for card in study_cards if card.get('category', 'General') == selected_category]

            if study_cards:
                # Initialize study state if not already
                if 'study_index' not in st.session_state:
                    st.session_state.study_index = 0
                    st.session_state.show_answer = False
                    st.session_state.cards_studied = 0
                    st.session_state.cards_correct = 0

                current_card = study_cards[st.session_state.study_index]

                # Progress bar
                progress = (st.session_state.study_index + 1) / len(study_cards)
                st.progress(progress, text=f"Card {st.session_state.study_index + 1} of {len(study_cards)}")

                # Front of card using css
                st.markdown(f"""
                <div style="
                    border: 2px solid #ddd;
                    border-radius: 10px;
                    padding: 30px;
                    margin: 20px 0;
                    background-color: #f9f9f9;
                    color: black;
                    text-align: center;
                    min-height: 150px;
                ">
                    <h3>{current_card['front']}</h3>
                </div>
                """, unsafe_allow_html=True)

                # Show answer + grading buttons
                if st.session_state.show_answer:
                    st.markdown(f"""
                    <div style="border: 2px solid #4CAF50; border-radius: 10px; padding: 20px; margin: 20px 0; 
                               background-color: #e8f5e8; color: black; text-align: center;">
                        <h4>Answer:</h4>
                        <p>{current_card['back']}</p>
                    </div>
                    """, unsafe_allow_html=True)

                    st.markdown("### How well did you know this?")
                    c1, c2, c3 = st.columns(3)

                    with c1:
                        if st.button("❌ Needs work", use_container_width=True):
                            user_data.save_current_user(st.session_state)
                            next_flashcard(study_cards, correct=False)

                    with c2:
                        if st.button("🤔 Almost", use_container_width=True):
                            user_data.save_current_user(st.session_state)
                            next_flashcard(study_cards, correct=True)

                    with c3:
                        if st.button("✅ Mastered", use_container_width=True):
                            user_data.save_current_user(st.session_state)
                            next_flashcard(study_cards, correct=True)
                else:
                    if st.button("🔍 Show Answer", use_container_width=True):
                        st.session_state.show_answer = True
                        st.rerun()

                # Session accuracy 
                if st.session_state.cards_studied > 0:
                    accuracy = (st.session_state.cards_correct / st.session_state.cards_studied) * 100
                    st.metric("Session Accuracy", f"{accuracy:.1f}%")

    # Create Tab
    with tab2:
        st.subheader("➕ Create Flashcards")

        method = st.radio(
            "Creation method:",
            ["📝 From Text", "📂 Upload File", "✋ Manual Entry", "📚 From Notes"],
            horizontal=True
        )

        # From Text
        if method == "📝 From Text":
            content = st.text_area("Paste content:", placeholder="Enter study material...", height=150)
            c1, c2, c3 = st.columns(3)
            with c1:
                num_cards = st.slider("Number of cards:", 3, 20, 8)
            with c2:
                difficulty = st.selectbox("Difficulty:", ["Easy", "Medium", "Hard"])
            with c3:
                category = st.text_input("Category:", value="General")

            if st.button("🚀 Generate Flashcards", type="primary"):
                if content.strip():
                    with st.spinner("Creating flashcards..."):
                        try:
                            flashcards = generators['flashcards'].generate_flashcards(
                                content, num_cards=num_cards, difficulty=difficulty
                            )
                            for card in flashcards:
                                card['category'] = category

                            st.session_state.flashcards.extend(flashcards)
                            auto_save()
                            st.success(f"✅ Generated {len(flashcards)} flashcards!")

                            # Log activity
                            session = {
                                'timestamp': datetime.now().isoformat(),
                                'activity_type': 'flashcards_created',
                                'subject': category,
                                'flashcards_created': len(flashcards)
                            }
                            st.session_state.study_sessions.append(session)
                            auto_save()

                            # Preview some cards
                            st.markdown("### Preview:")
                            for i, card in enumerate(flashcards[:3], 1):
                                with st.expander(f"Card {i}"):
                                    st.write(f"**Front:** {card['front']}")
                                    st.write(f"**Back:** {card['back']}")
                            if len(flashcards) > 3:
                                st.info(f"+ {len(flashcards) - 3} more cards created!")
                                ser_data.save_current_user(st.session_state)
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
                else:
                    st.warning("Please enter content.")

        # Manual Entry
        elif method == "✋ Manual Entry":
            with st.form("manual_flashcard"):
                front = st.text_area("Front (Question):", height=100)
                back = st.text_area("Back (Answer):", height=100)
                category = st.text_input("Category:", value="General")

                if st.form_submit_button("➕ Add Flashcard"):
                    user_data.save_current_user(st.session_state)
                    if front.strip() and back.strip():
                        new_card = {
                            'front': front,
                            'back': back,
                            'category': category,
                            'created': datetime.now().isoformat()
                        }
                        st.session_state.flashcards.append(new_card)
                        auto_save()
                        user_data.save_current_user(st.session_state)
                        st.success("✅ Flashcard added!")
                    else:
                        st.warning("Please fill in both sides.")

        # Upload File
        elif method == "📂 Upload File":
            uploaded_file = st.file_uploader("Choose a file", type=["txt", "pdf", "docx"])
            content = ""
            if uploaded_file is not None:
                if uploaded_file.type == "text/plain":
                    content = uploaded_file.read().decode("utf-8")
                elif uploaded_file.type == "application/pdf":
                    pdf = PdfReader(uploaded_file)
                    content = "".join(page.extract_text() or "" for page in pdf.pages)
                elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    doc = docx.Document(uploaded_file)
                    content = "\n".join([p.text for p in doc.paragraphs])
                st.text_area("Preview:", value=(content[:200] + "...") if content else "", height=100, disabled=True)

            c1, c2, c3 = st.columns(3)
            with c1:
                num_cards = st.slider("Number of cards:", 3, 20, 8)
            with c2:
                difficulty = st.selectbox("Difficulty:", ["Easy", "Medium", "Hard"])
            with c3:
                category = st.text_input("Category:", value="General")
            # Generate from file content
            if st.button("🚀 Generate Flashcards", type="primary"):
                user_data.save_current_user(st.session_state)
                if content.strip():
                    with st.spinner("Creating flashcards..."):
                        try:
                            flashcards = generators['flashcards'].generate_flashcards(
                                content, num_cards=num_cards, difficulty=difficulty
                            )
                            for card in flashcards:
                                card['category'] = category

                            st.session_state.flashcards.extend(flashcards)
                            auto_save()
                            user_data.save_current_user(st.session_state)
                            st.success(f"✅ Generated {len(flashcards)} flashcards!")

                            # Log activity
                            session = {
                                'timestamp': datetime.now().isoformat(),
                                'activity_type': 'flashcards_created',
                                'subject': category,
                                'flashcards_created': len(flashcards)
                            }
                            st.session_state.study_sessions.append(session)
                            auto_save()

                            # Preview some cards
                            st.markdown("### Preview:")
                            for i, card in enumerate(flashcards[:3], 1):
                                with st.expander(f"Card {i}"):
                                    st.write(f"**Front:** {card['front']}")
                                    st.write(f"**Back:** {card['back']}")
                            if len(flashcards) > 3:
                                st.info(f"+ {len(flashcards) - 3} more cards created!")
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
                else:
                    st.warning("Please enter content.")

        # From Notes
        elif method == "📚 From Notes":
            if st.session_state.notes:
                note_titles = [n['title'] for n in st.session_state.notes]
                selected_note = st.selectbox("Select note:", note_titles)
                note_obj = next((n for n in st.session_state.notes if n['title'] == selected_note), None)
                if note_obj:
                    content = note_obj['content']
                    st.text_area("Preview:", value=content[:200] + "...", height=100, disabled=True)

                    c1, c2, c3 = st.columns(3)
                    with c1:
                        num_cards = st.slider("Number of cards:", 3, 20, 8)
                    with c2:
                        difficulty = st.selectbox("Difficulty:", ["Easy", "Medium", "Hard"])
                    with c3:
                        category = st.text_input("Category:", value="General")
                    # generate from note content
                    if st.button("🚀 Generate Flashcards from Note", type="primary"):
                        user_data.save_current_user(st.session_state)
                        if content.strip():
                            with st.spinner("Creating flashcards..."):
                                try:
                                    flashcards = generators['flashcards'].generate_flashcards(
                                        content, num_cards=num_cards, difficulty=difficulty
                                    )
                                    for card in flashcards:
                                        card['category'] = category

                                    st.session_state.flashcards.extend(flashcards)
                                    auto_save()
                                    user_data.save_current_user(st.session_state)
                                    st.success(f"✅ Generated {len(flashcards)} flashcards from note!")

                                    # Log activity
                                    session = {
                                        'timestamp': datetime.now().isoformat(),
                                        'activity_type': 'flashcards_created',
                                        'subject': category,
                                        'flashcards_created': len(flashcards)
                                    }
                                    st.session_state.study_sessions.append(session)
                                    auto_save()

                                    # Preview
                                    st.markdown("### Preview:")
                                    for i, card in enumerate(flashcards[:3], 1):
                                        with st.expander(f"Card {i}"):
                                            st.write(f"**Front:** {card['front']}")
                                            st.write(f"**Back:** {card['back']}")
                                    if len(flashcards) > 3:
                                        st.info(f"+ {len(flashcards) - 3} more cards created!")
                                except Exception as e:
                                    st.error(f"Error: {str(e)}")
                        else:
                            st.warning("Note is empty.")
            else:
                st.info("No notes available. Create some first!")

    # Manage Tab
    with tab3:
        st.subheader("📂 Manage Flashcards")

        if st.session_state.flashcards:
            st.write(f"**Total flashcards:** {len(st.session_state.flashcards)}")

            # Export and clear all
            c1, c2 = st.columns(2)
            with c1:
                if st.button("📥 Export All"):
                    data = generators['flashcards'].save_flashcards_file(
                        st.session_state.flashcards, "export"
                    )
                    st.download_button(
                        "Download",
                        data=data,
                        file_name=f"flashcards_{datetime.now().strftime('%Y%m%d')}.json",
                        mime="application/json"
                    )
            with c2:
                if st.button("🗑️ Clear All"):
                    user_data.save_current_user(st.session_state)
                    st.session_state.flashcards = []
                    auto_save()
                    st.success("✅ All flashcards deleted!")
                    st.rerun()

            # Filter and display
            categories = list(set([card.get('category', 'General') for card in st.session_state.flashcards]))
            filter_cat = st.selectbox("Filter:", ["All"] + categories)

            filtered = st.session_state.flashcards
            if filter_cat != "All":
                filtered = [c for c in filtered if c.get('category', 'General') == filter_cat]

            for i, card in enumerate(filtered):
                with st.expander(f"🎴 {card['front'][:50]}..."):
                    st.write(f"**Front:** {card['front']}")
                    st.write(f"**Back:** {card['back']}")
                    st.write(f"**Category:** {card.get('category', 'General')}")

                    if st.button("🗑️ Delete", key=f"del_{i}"):
                        user_data.save_current_user(st.session_state)
                        st.session_state.flashcards.remove(card)
                        auto_save()
                        st.rerun()
        else:
            st.info("No flashcards yet. Create some first!")


# ============================
# Quizzes Page
# ============================

elif st.session_state.page == "🧠 Quizzes":
    st.title("🧠 Interactive Quiz System")

    tab1, tab2 = st.tabs(["📝 Take Quiz", "📊 History"])

    # Take Quiz
    with tab1:
        # If a quiz is active, display it with AdvancedQuizSystem.py
        if st.session_state.get('quiz_active', False):
            advanced_quiz.display_quiz_interface(st.session_state.get('current_quiz'))
        else:
            st.subheader("📝 Create New Quiz")

            content = ""  # Will hold the text used to generate the quiz

            c1, c2 = st.columns(2)
            with c1:
                source = st.radio("Quiz source:", ["📚 My Notes", "📝 New Content", "📂 Upload file"])
                question_type = st.selectbox(
                    "Question Type:",
                    options=[
                        "Multiple Choice Only",
                        "True/False Only",
                        "Short Answer Only",
                        "Mixed Questions"
                    ],
                    index=3
                )
            with c2:
                num_questions = st.slider("Questions:", 3, 15, 8)
                difficulty = st.selectbox("Difficulty:", ["Easy", "Medium", "Hard"])

            # Choose content source
            if source == "📚 My Notes":
                if st.session_state.notes:
                    note_titles = [n['title'] for n in st.session_state.notes]
                    selected_note = st.selectbox("Select note:", note_titles)
                    note_obj = next((n for n in st.session_state.notes if n['title'] == selected_note), None)
                    if note_obj:
                        content = note_obj['content']
                        st.text_area("Preview:", value=content[:200] + "...", height=100, disabled=True)
                else:
                    st.info("No notes available. Create some first!")

            elif source == "📝 New Content":
                content = st.text_area("Enter content:", height=200, placeholder="Paste study material...")

            elif source == "📂 Upload file":
                uploaded_file = st.file_uploader("Choose a file", type=["txt", "pdf", "docx"])
                if uploaded_file is not None:
                    if uploaded_file.type == "text/plain":
                        content = uploaded_file.read().decode("utf-8")
                    elif uploaded_file.type == "application/pdf":
                        pdf = PdfReader(uploaded_file)
                        content = "".join(page.extract_text() or "" for page in pdf.pages)
                    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        doc = docx.Document(uploaded_file)
                        content = "\n".join([p.text for p in doc.paragraphs])
                    st.text_area("Preview:", value=content[:200] + "...", height=100, disabled=True)

            # Start quiz
            if st.button("🚀 Create & Start Quiz", type="primary", use_container_width=True):
                user_data.save_current_user(st.session_state)
                if content.strip():
                    with st.spinner("Creating quiz..."):
                        try:
                            quiz_data = advanced_quiz.create_quiz_from_content(
                                content,
                                num_questions=num_questions,
                                difficulty=difficulty,
                                question_type=question_type
                            )
                            if quiz_data and quiz_data.get('questions'):
                                st.session_state.current_quiz = quiz_data
                                st.session_state.quiz_active = True
                                st.session_state.quiz_question_index = 0
                                st.session_state.quiz_answers = {}
                                st.success("✅ Quiz created! Starting now...")
                                st.rerun()
                            else:
                                st.error("Failed to create quiz. Please try again.")
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
                else:
                    st.warning("Please provide content for the quiz.")

    # History
    with tab2:
        st.subheader("📊 Quiz History")

        quiz_sessions = [s for s in st.session_state.study_sessions if s.get('activity_type') == 'quiz']

        if quiz_sessions:
            c1, c2, c3 = st.columns(3)
            with c1:
                st.metric("Total Quizzes", len(quiz_sessions))
            with c2:
                avg_score = sum(s.get('score', 0) for s in quiz_sessions) / len(quiz_sessions)
                st.metric("Average Score", f"{avg_score:.1f}%")
            with c3:
                best_score = max(s.get('score', 0) for s in quiz_sessions)
                st.metric("Best Score", f"{best_score:.1f}%")

            st.subheader("Recent Results")
            for i, session in enumerate(reversed(quiz_sessions[-10:])):
                timestamp = datetime.fromisoformat(session['timestamp']).strftime("%Y-%m-%d %H:%M")
                score = session.get('score', 0)
                correct = session.get('correct_answers', 0)
                total = session.get('total_questions', 0)
                color = "🟢" if score >= 80 else "🟡" if score >= 60 else "🔴"

                with st.expander(f"{color} {timestamp} - {score:.1f}% ({correct}/{total})"):
                    st.write(f"**Score:** {score:.1f}%")
                    st.write(f"**Difficulty:** {session.get('difficulty', 'N/A')}")

                    # Retake
                    if st.button("🔄 Retake This Quiz", key=f"retake_{i}"):
                        user_data.save_current_user(st.session_state)
                        st.session_state.retake_quiz_content = session.get('original_content', '')
                        st.session_state.retake_quiz_config = {
                            'num_questions': session.get('total_questions', 10),
                            'difficulty': session.get('difficulty', 'Medium')
                        }
                        st.session_state.page = "🧠 Quizzes"
                        st.session_state.quiz_active = False
                        st.rerun()
        else:
            st.info("No quiz history yet. Take your first quiz!")


# ============================
# Progress Page
# ============================

elif st.session_state.page == "📊 Progress":
    st.title("📊 Learning Analytics")

    if not st.session_state.study_sessions:
        st.info("No study data yet. Start using the platform to see your progress!")
    else:
        # Summary metrics
        c1, c2, c3, c4 = st.columns(4)

        total_sessions = len(st.session_state.study_sessions)
        quiz_sessions = [s for s in st.session_state.study_sessions if s.get('activity_type') == 'quiz']
        flashcard_sessions = [s for s in st.session_state.study_sessions if s.get('activity_type') in ['flashcards', 'flashcards_created']]

        with c1:
            st.metric("Total Sessions", total_sessions)
        with c2:
            st.metric("Quizzes Taken", len(quiz_sessions))
        with c3:
            if quiz_sessions:
                avg_quiz_score = sum(s.get('score', 0) for s in quiz_sessions) / len(quiz_sessions)
                st.metric("Avg Quiz Score", f"{avg_quiz_score:.1f}%")
            else:
                st.metric("Avg Quiz Score", "N/A")
        with c4:
            st.metric("Flashcard Sessions", len(flashcard_sessions))

        # Recent activity bar chart (last 7 days)
        st.subheader("📈 Recent Activity")

        daily_activity = defaultdict(int)
        for session in st.session_state.study_sessions:
            try:
                date = datetime.fromisoformat(session['timestamp']).strftime('%Y-%m-%d')
                daily_activity[date] += 1
            except Exception:
                # Skip malformed timestamps
                continue

        if daily_activity:
            dates = list(daily_activity.keys())[-7:]
            counts = [daily_activity[date] for date in dates]

            fig, ax = plt.subplots(figsize=(10, 4))
            ax.bar(dates, counts)
            ax.set_title('Study Sessions (Last 7 Days)')
            ax.set_ylabel('Sessions')
            plt.xticks(rotation=45)
            st.pyplot(fig)

# ============================
# Calendar Page
# ============================

elif st.session_state.page == "📅 Calendar":
    st.title("📅 Calendar & Events")


    # Initialize calendar sessions
    if "events" not in st.session_state:
        st.session_state.events = []
    if "calendar_year" not in st.session_state:
        st.session_state.calendar_year = datetime.now().year
    if "calendar_month" not in st.session_state:
        st.session_state.calendar_month = datetime.now().month
    if "selected_date" not in st.session_state:
        st.session_state.selected_date = None

    # Cache calendar grid for instant reruns
    @st.cache_data
    def build_calendar_html(events, year, month):
        import calendar
        from datetime import datetime
        cal = calendar.Calendar(firstweekday=0)
        month_days = list(cal.itermonthdates(year, month))
        html = """
        <style>
            .calendar-day {padding: 10px; background: #f9f9f9; border-radius: 8px; border: 1px solid #ddd; min-height: 80px; text-align: left; position: relative; transition: all 0.2s ease;}
            .calendar-day:hover {background: #f0f0f0; transform: translateY(-2px); box-shadow: 0 2px 5px rgba(0,0,0,0.1);}
            .day-number {position: absolute; top: 5px; right: 8px; font-weight: 600; font-size: 14px; color: #333;}
            .today-highlight {border: 2px solid #2196F3 !important; background: #e3f2fd !important;}
            .event-item {margin: 3px 0; border-radius: 4px; font-size: 11px; padding: 3px 5px; white-space: nowrap; overflow: hidden; text-overflow: ellipsis; box-shadow: 0 1px 2px rgba(0,0,0,0.1);}
            .event-notes {font-size: 10px; color: #555; margin-top: 2px; white-space: normal;}
            .more-events {font-size: 10px; color: #666; margin-top: 3px; font-style: italic;}
        </style>
        <div style='display:grid; grid-template-columns: repeat(7, 1fr); gap:8px; font-family:sans-serif; text-align:center; width:100%;'>
        """
        for wd in ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]:
            html += f"<div style='font-weight:bold; padding:10px; background:#f0f0f0; border-radius:4px; color: black;'>{wd}</div>"
        today = datetime.now().date()
        for day in month_days:
            if day.month == month:
                events_today = [e for e in events if datetime.fromisoformat(e["date"]).date() == day]
                today_class = "today-highlight" if day == today else ""
                html += f"<div class='calendar-day {today_class}'>"
                html += f"<div class='day-number'>{day.day}</div>"
                if events_today:
                    html += "<div style='margin-top:20px; max-height:60px; overflow-y:auto; padding-right:3px;'>"
                    for e in events_today[:4]:
                        html += f"<div class='event-item' style='background:{e['color']}; color:white;' title='{e['name']}'>"
                        html += f"{e['name']}"
                        if e.get('notes'):
                            html += f"<div class='event-notes'>{e['notes']}</div>"
                        html += "</div>"
                    if len(events_today) > 4:
                        html += f"<div class='more-events'>+{len(events_today)-4} more</div>"
                    html += "</div>"
                html += "</div>"
            else:
                html += "<div style='padding:10px; color:#ccc; min-height:80px;'></div>"
        html += "</div>"
        return html

    # Month navigation helper (basically tracks the date and changes the year)
    def change_month(delta):
        new_month = st.session_state.calendar_month + delta
        if new_month > 12:
            st.session_state.calendar_month = 1
            st.session_state.calendar_year += 1
        elif new_month < 1:
            st.session_state.calendar_month = 12
            st.session_state.calendar_year -= 1
        else:
            st.session_state.calendar_month = new_month

    # Add Event form
    st.subheader("➕ Add Event")
    with st.form("add_event_form"):
        user_data.save_current_user(st.session_state)
        name = st.text_input("Event Title:", placeholder="e.g., Math Test, History Project, Concert")
        date = st.date_input("Date:")
        notes = st.text_area("Details (optional):", placeholder="Extra info...")
        color = st.color_picker("Pick a color:", "#4CAF50")  # default green

        submitted = st.form_submit_button("➕ Add")
        if submitted:
            if name.strip():
                new_event = {
                    "name": name,
                    "date": date.isoformat(),
                    "notes": notes,
                    "color": color,
                    "created": datetime.now().isoformat()
                }
                st.session_state.events.append(new_event)
                user_data.save_current_user(st.session_state)
                auto_save()
                st.success(f"✅ Added event - {name}")
            else:
                st.warning("Please enter a title.")

    st.divider()

    # Month Navigation Controls
    c1, c2, c3 = st.columns([1, 3, 1])
    with c1:
        st.button("‹", key="prev_month", on_click=change_month, args=(-1,))
    with c2:
        month_name = datetime(
            st.session_state.calendar_year,
            st.session_state.calendar_month,
            1
        ).strftime("%B %Y")
        st.markdown(f"<h2 style='text-align:center;margin:0'>{month_name}</h2>", unsafe_allow_html=True)
    with c3:
        st.button("›", key="next_month", on_click=change_month, args=(1,))

    st.divider()

    # Build the month grid
    cal = calendar.Calendar(firstweekday=0)  # Monday start (0=Monday)
    month_days = list(cal.itermonthdates(st.session_state.calendar_year, st.session_state.calendar_month))

    # Inline CSS and HTML grid layout
    html = """
    <style>
        .calendar-day {
            padding: 10px;
            background: #f9f9f9;
            border-radius: 8px;
            border: 1px solid #ddd;
            min-height: 80px;
            text-align: left;
            position: relative;
            transition: all 0.2s ease;
        }
        .calendar-day:hover {
            background: #f0f0f0;
            transform: translateY(-2px);
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }
        .day-number {
            position: absolute;
            top: 5px;
            right: 8px;
            font-weight: 600;
            font-size: 14px;
            color: #333;
        }
        .today-highlight {
            border: 2px solid #2196F3 !important;
            background: #e3f2fd !important;
        }
        .event-item {
            margin: 3px 0;
            border-radius: 4px;
            font-size: 11px;
            padding: 3px 5px;
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
            box-shadow: 0 1px 2px rgba(0,0,0,0.1);
        }
        .event-notes {
            font-size: 10px;
            color: #555;
            margin-top: 2px;
            white-space: normal;
        }
        .more-events {
            font-size: 10px;
            color: #666;
            margin-top: 3px;
            font-style: italic;
        }
    </style>
    <div style='display:grid; grid-template-columns: repeat(7, 1fr); gap:8px; font-family:sans-serif; text-align:center; width:100%;'>
    """

    # Weekday headers
    for wd in ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]:
        html += f"<div style='font-weight:bold; padding:10px; background:#f0f0f0; border-radius:4px; color: black;'>{wd}</div>"

    today = datetime.now().date()

    # Render each day cell with events
    for day in month_days:
        if day.month == st.session_state.calendar_month:
            events_today = [e for e in st.session_state.events if datetime.fromisoformat(e["date"]).date() == day]

            today_class = "today-highlight" if day == today else ""

            html += f"<div class='calendar-day {today_class}'>"
            html += f"<div class='day-number'>{day.day}</div>"

            if events_today:
                html += "<div style='margin-top:20px; max-height:60px; overflow-y:auto; padding-right:3px;'>"
                for e in events_today[:4]:
                    html += f"<div class='event-item' style='background:{e['color']}; color:white;' title='{e['name']}'>"
                    html += f"{e['name']}"
                    if e.get('notes'):
                        html += f"<div class='event-notes'>{e['notes']}</div>"
                    html += "</div>"
                if len(events_today) > 4:
                    html += f"<div class='more-events'>+{len(events_today)-4} more</div>"
                html += "</div>"

            html += "</div>"
        else:
            # Blank cell for days outside current month grid
            html += "<div style='padding:10px; color:#ccc; min-height:80px;'></div>"

    html += "</div>"

    # Display Calendar grid
    st.markdown(html, unsafe_allow_html=True)

    # Today's reminders
    st.divider()
    st.subheader("📅 Today's Reminders")
    today_events = [e for e in st.session_state.events if datetime.fromisoformat(e["date"]).date() == today]
    if today_events:
        for e in today_events:
            with st.container(border=True):
                st.markdown(f"<span style='color:{e['color']};font-size:20px'>●</span> **{e['name']}**", unsafe_allow_html=True)
                if e.get('notes'):
                    st.caption(f"📝 {e['notes']}")
    else:
        st.info("No events today.")

    # Upcoming reminders
    st.divider()
    st.subheader("⏰ Upcoming")
    upcoming = [e for e in st.session_state.events if datetime.fromisoformat(e["date"]).date() > today]
    upcoming = sorted(upcoming, key=lambda x: x["date"])[:5]
    if upcoming:
        for e in upcoming:
            with st.container(border=True):
                date_obj = datetime.fromisoformat(e["date"]).strftime("%a, %b %d")
                st.markdown(f"<span style='color:{e['color']};font-size:20px'>●</span> **{date_obj}** — {e['name']}", unsafe_allow_html=True)
                if e.get('notes'):
                    st.caption(f"📝 {e['notes']}")
    else:
        st.info("No upcoming events!")

    # Delete event UI
    st.divider()
    st.subheader("🗑️ Delete Event")

    if st.session_state.events:
        # Dropdown to select date
        event_options = [
            f"{datetime.fromisoformat(e['date']).strftime('%Y/%m/%d')} — {e['name']}"
            for e in st.session_state.events
        ]
        event_to_delete = st.selectbox("Select an event to delete:", options=event_options)

        if st.button("❌ Delete Selected Event"):
            user_data.save_current_user(st.session_state)
            for e in list(st.session_state.events):
                label = f"{datetime.fromisoformat(e['date']).strftime('%Y/%m/%d')} — {e['name']}"
                if label == event_to_delete:
                    st.session_state.events.remove(e)
                    auto_save()
                    st.success(f"Deleted event: {label}")
                    st.rerun()
                    user_data.save_current_user(st.session_state)
    else:
        st.info("No events to delete.")
# ============================
# Autograder Page
# ============================
elif st.session_state.page == "📝 Autograder":
    from autograder import AutoGrader
    st.title("📝 AI Autograder")
    # Input Section

    text_input = st.text_area("✍️ Paste your essay, story, or text:", height=250, placeholder="Write or paste your text here...")

    col1, col2 = st.columns(2)
    with col1:
        text_type = st.selectbox("Text type:", ["Essay", "Story", "Article", "Other"])
    with col2:
        extra_notes = st.text_input("Extra notes (optional)", placeholder="e.g., Focus on creativity, academic tone...")

    if st.button("🚀 Grade Now", type="primary", use_container_width=True):
        if not text_input.strip():
            st.warning("Please enter some text to grade.")
        else:
            grader = AutoGrader()
            with st.spinner("Grading with AI..."):
                result = grader.grade_text(text_input, text_type, extra_notes)

            # Stylish Results
            st.subheader(f"📊 Score: {result.get('score', 0)}/10")
            st.progress(int(result.get("score", 0)) / 10)

            st.markdown("### ✅ Strengths")
            for s in result.get("strengths", []):
                st.markdown(f"- {s}")

            st.markdown("### ⚠️ Weaknesses")
            for w in result.get("weaknesses", []):
                st.markdown(f"- {w}")

            st.markdown("### 💡 Suggestions to Improve")
            for sug in result.get("suggestions", []):
                st.markdown(f"- {sug}")

            st.markdown("### 📝 Detailed Feedback")
            st.info(result.get("detailed_feedback", "No feedback provided."))
# ============================
# Settings Page
# ============================

elif st.session_state.page == "⚙️ Settings":
    st.title("⚙️ Account Settings")
    
    # Change Password Section
    st.subheader("🔐 Change Password")
    
    with st.form("change_password_form"):
        current_password = st.text_input("Current Password", type="password", 
                                       placeholder="Enter your current password")
        new_password = st.text_input("New Password", type="password", 
                                   placeholder="Enter new password")
        confirm_password = st.text_input("Confirm New Password", type="password", 
                                       placeholder="Re-enter new password")
        
        if st.form_submit_button("Change Password", use_container_width=True):
            if not current_password or not new_password or not confirm_password:
                st.error("Please fill in all password fields.")
            elif new_password != confirm_password:
                st.error("New passwords do not match.")
            elif len(new_password) < 6:
                st.error("Password must be at least 6 characters long.")
            else:
                # Verify current password first
                ok, msg = user_data.authenticate(st.session_state["username"], current_password)
                if ok:
                    # Change the password
                    success, message = user_data.change_password(st.session_state["username"], new_password)
                    if success:
                        st.success("✅ Password changed successfully!")
                    else:
                        st.error(f"Failed to change password: {message}")
                else:
                    st.error("Current password is incorrect.")
    
    st.markdown("---")
    
    # Data Management Section
    st.subheader("📊 Data Management")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📥 Export All Data", use_container_width=True):
            with st.spinner("Exporting your data..."):
                success, message, exported_data = user_data.export_user_data(st.session_state["username"])
                
                if success:
                    import json
                    from datetime import datetime
                    
                    # Create filename with timestamp
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"study_platform_export_{st.session_state['username']}_{timestamp}.json"
                    
                    # Convert to JSON string
                    json_data = json.dumps(exported_data, indent=2, ensure_ascii=False)
                    
                    # Show export summary
                    st.success("✅ Data exported successfully!")
                    
                    # Display export summary
                    meta = exported_data.get("export_metadata", {})
                    st.info(f"""
                    **Export Summary:**
                    - 📝 Notes: {len(exported_data.get('notes', []))}
                    - 🎴 Flashcards: {len(exported_data.get('flashcards', []))}
                    - 📚 Study Sessions: {len(exported_data.get('study_sessions', []))}
                    - 📅 Events: {len(exported_data.get('events', []))}
                    - 📦 Total Items: {meta.get('total_items', 0)}
                    """)
                    
                    # Download button
                    st.download_button(
                        label="⬇️ Download JSON File",
                        data=json_data,
                        file_name=filename,
                        mime="application/json",
                        use_container_width=True
                    )
                    
                    # Show preview
                    with st.expander("🔍 Preview Exported Data"):
                        st.json(exported_data)
                else:
                    st.error(f"Export failed: {message}")
    
    with col2:
        if st.button("🔄 Refresh Data", use_container_width=True):
            # Reload data from database
            ok, data = user_data.load_user_data(
                st.session_state["username"], 
                merge_local=False, 
                local_state=st.session_state
            )
            if ok:
                st.session_state.update(data)
                st.success("✅ Data refreshed from server!")
            else:
                st.error("Failed to refresh data.")
    
    st.markdown("---")
    
    # Delete Account Section
    st.subheader("🗑️ Delete Account")
    
    st.warning("⚠️ **Danger Zone**: This action cannot be undone! All your data will be permanently deleted.")
    
    with st.form("delete_account_form"):
        confirm_username = st.text_input("Confirm your username to delete account", 
                                       placeholder="Type your username to confirm")
        confirm_password = st.text_input("Confirm your password", type="password", 
                                       placeholder="Enter your password to confirm")
        
        if st.form_submit_button("🚨 Delete My Account", type="secondary", use_container_width=True):
            if not confirm_username or not confirm_password:
                st.error("Please enter both username and password to confirm deletion.")
            elif confirm_username != st.session_state["username"]:
                st.error("Username does not match your account.")
            else:
                # Verify password first
                ok, msg = user_data.authenticate(st.session_state["username"], confirm_password)
                if ok:
                    # Delete the account
                    success, message = user_data.delete_account(st.session_state["username"])
                    if success:
                        st.success("✅ Account deleted successfully!")
                        st.info("You will be logged out automatically.")
                        # Clear session and log out
                        st.session_state.clear()
                        st.session_state["logged_in"] = False
                        st.session_state["page"] = "🏠 Home"
                        st.rerun()
                    else:
                        st.error(f"Failed to delete account: {message}")
                else:
                    st.error("Password is incorrect.")
    
    st.markdown("---")
    
    # Account Information
    st.subheader("📋 Account Information")
    
    info_col1, info_col2 = st.columns(2)
    
    with info_col1:
        st.metric("Username", st.session_state["username"])
        st.metric("Notes Created", len(st.session_state.get("notes", [])))
    
    with info_col2:
        st.metric("Flashcards", len(st.session_state.get("flashcards", [])))
        st.metric("Study Sessions", len(st.session_state.get("study_sessions", [])))